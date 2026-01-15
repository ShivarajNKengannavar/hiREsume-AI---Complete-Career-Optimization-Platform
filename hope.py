# --- Required Imports ---
import streamlit as st
from streamlit_lottie import st_lottie
from streamlit_extras.add_vertical_space import add_vertical_space
from streamlit_extras.annotated_text import annotated_text
import google.generativeai as genai
from google.generativeai.types import GenerationConfig
import os
import PyPDF2 as pdf
from dotenv import load_dotenv
import json
import io
import pandas as pd
import base64
import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH # Required for docx alignment
import docx.oxml.ns # Required for DOCX shading
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
import datetime
import re
import zipfile
from reportlab.lib.pagesizes import A4, letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY, TA_RIGHT
from reportlab.lib.colors import HexColor 
from reportlab.platypus import Table, TableStyle
from reportlab.lib import colors
import requests
import math

# Import PDF styles from pdf_styles
from pdf_styles import global_styles, table_styles

# Use the default table style
table_style = table_styles['default']

# Get the styles from global_styles
styles = global_styles

# Initialize session state variables
if 'job_results' not in st.session_state:
    st.session_state.job_results = None


def get_docx_styles(doc):
    # Add custom styles to the document
    styles = doc.styles
    
    # Normal style (justified)
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(12)
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Title style
    title_style = styles['Title']
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    title_style.paragraph_format.space_after = Pt(12)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Heading style
    heading_style = styles['Heading 1']
    heading_style.font.name = 'Calibri'
    heading_style.font.size = Pt(14)
    heading_style.font.bold = True
    heading_style.paragraph_format.space_after = Pt(6)
    heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    return styles

def export_to_pdf(jobs, filename="job_search_results.pdf"):
    doc = SimpleDocTemplate(
        filename,
        pagesize=letter,
        rightMargin=40, 
        leftMargin=40,
        topMargin=40, 
        bottomMargin=40
    )
    
    # Using imported styles from pdf_styles.py
    elements = []
    
    # Title
    elements.append(Paragraph("Job Search Results", styles['Title']))
    elements.append(Spacer(1, 12))
    
    # Add each job
    for i, job in enumerate(jobs, 1):
        # Job title
        elements.append(Paragraph(f"{i}. {job.get('job_title', 'No Title')}", styles['Heading2']))
        
        # Company and location
        company = job.get('employer_name', 'Unknown Company')
        location = job.get('job_city', '') + (f", {job.get('job_state', '')}" if job.get('job_state') else '')
        elements.append(Paragraph(f"<b>Company:</b> {company} | <b>Location:</b> {location}", styles['Normal']))
        
        # Job type and salary
        job_type = job.get('job_employment_type', 'Not specified')
        salary = job.get('job_salary', 'Salary not provided')
        elements.append(Paragraph(f"<b>Type:</b> {job_type} | <b>Salary:</b> {salary}", styles['Justify']))
        
        # Job description
        description = job.get('job_description', 'No description available')
        elements.append(Paragraph("<b>Description:</b>", styles['Justify']))
        elements.append(Paragraph(description, styles['Justify']))
        
        elements.append(Spacer(1, 12))
    
    # Build the PDF
    doc.build(elements)
    return filename

def export_to_docx(jobs, filename="job_search_results.docx"):
    doc = Document()
    styles = get_docx_styles(doc)
    
    # Title
    doc.add_paragraph('Job Search Results', style='Title')
    
    # Add each job
    for i, job in enumerate(jobs, 1):
        # Job title
        doc.add_paragraph(f"{i}. {job.get('job_title', 'No Title')}", style='Heading 1')
        
        # Company and location
        company = job.get('employer_name', 'Unknown Company')
        location = job.get('job_city', '') + (f", {job.get('job_state', '')}" if job.get('job_state') else '')
        doc.add_paragraph(f"Company: {company} | Location: {location}")
        
        # Job type and salary
        job_type = job.get('job_employment_type', 'Not specified')
        salary = job.get('job_salary', 'Salary not provided')
        doc.add_paragraph(f"Type: {job_type} | Salary: {salary}")
        
        # Job description
        description = job.get('job_description', 'No description available')
        doc.add_paragraph("Description:", style='Heading 2')
        doc.add_paragraph(description)
        
        doc.add_paragraph()  # Add empty paragraph for spacing
    
    doc.save(filename)
    return filename

def add_export_buttons(jobs):
    if not jobs:
        return
        
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📄 Export to PDF", use_container_width=True):
            with st.spinner("Generating PDF..."):
                filename = export_to_pdf(jobs)
                with open(filename, "rb") as f:
                    st.download_button(
                        label="⬇️ Download PDF",
                        data=f,
                        file_name=filename,
                        mime="application/pdf",
                        use_container_width=True
                    )
                os.remove(filename)
    
    with col2:
        if st.button("📄 Export to DOCX", use_container_width=True):
            with st.spinner("Generating Word Document..."):
                filename = export_to_docx(jobs)
                with open(filename, "rb") as f:
                    st.download_button(
                        label="⬇️ Download Word",
                        data=f,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
def create_pdf_export(jobs, filename="job_search_results.pdf"):
    # Use A4 size with larger margins
    doc = SimpleDocTemplate(
        filename,
        pagesize=A4,
        rightMargin=36,  # 0.5 inch
        leftMargin=36,   # 0.5 inch
        topMargin=36,    # 0.5 inch
        bottomMargin=36  # 0.5 inch
    )
    
    styles = get_pdf_styles()
    elements = []
    
    # Title
    elements.append(Paragraph("Job Search Results", styles['Title']))
    elements.append(Spacer(1, 12))
    
    # Add each job
    for i, job in enumerate(jobs, 1):
        # Job title
        elements.append(Paragraph(f"{i}. {job.get('job_title', 'No Title')}", styles['Heading']))
        
        # Company and location
        company = job.get('employer_name', 'Unknown Company')
        location = job.get('job_city', '') + (f", {job.get('job_state', '')}" if job.get('job_state') else '')
        elements.append(Paragraph(f"<b>Company:</b> {company} | <b>Location:</b> {location}", styles['Justify']))
        
        # Job type and salary
        job_type = job.get('job_employment_type', 'Not specified')
        salary = job.get('job_salary', 'Salary not provided')
        elements.append(Paragraph(f"<b>Type:</b> {job_type} | <b>Salary:</b> {salary}", styles['Justify']))
        
        # Job description with better formatting
        description = job.get('job_description', 'No description available')
        # Split description into paragraphs and add proper spacing
        paragraphs = [p.strip() for p in description.split('\n') if p.strip()]
        elements.append(Paragraph("<b>Description:</b>", styles['Justify']))
        
        for para in paragraphs:
            # Truncate very long paragraphs to prevent overflow
            if len(para) > 1000:
                para = para[:1000] + "... [truncated]"
            elements.append(Paragraph(para, styles['JobDesc']))
            elements.append(Spacer(1, 4))  # Small space between paragraphs
        
        # Add page break if we're getting close to the end of the page
        if i < len(jobs):  # Don't add after last job
            elements.append(PageBreak())
    
    # Build the PDF with page numbers
    def add_page_number(canvas, doc):
        canvas.saveState()
        page_num = canvas.getPageNumber()
        text = f"Page {page_num}"
        canvas.setFont('Helvetica', 8)
        canvas.drawRightString(7.5*inch, 0.5*inch, text)
        canvas.restoreState()
    
    doc.build(elements, onFirstPage=add_page_number, onLaterPages=add_page_number)
    return filename
    # Build the PDF
    doc.build(elements)
    return filename

# Add this where you want the export button to appear (after displaying job results)
if st.session_state.job_results:
    if st.button("📄 Export to PDF"):
        with st.spinner("Generating PDF..."):
            filename = create_pdf_export(st.session_state.job_results)
            with open(filename, "rb") as f:
                st.download_button(
                    label="⬇️ Download PDF",
                    data=f,
                    file_name=filename,
                    mime="application/pdf"
                )
            # Clean up the temporary file
            os.remove(filename)

if "global_resume_text" not in st.session_state:
    st.session_state.global_resume_text = ""

if "global_resume_filename" not in st.session_state:
    st.session_state.global_resume_filename = ""

# ---------------------------
# JOB SEARCH (Advanced) Block
# Paste/replace your existing "🔍 Job Search" section with this block
# ---------------------------

# Helper: map experience to API tokens (copy into a helpers section if you already have it)
def map_experience_to_api(exp_range: str):
    mapping = {
        "Any": "",
        "Fresher (0 years)": "intern",
        "0–1 years": "entry",
        "1–3 years": "entry",
        "3–5 years": "mid",
        "5–8 years": "mid",
        "8+ years": "senior"
    }
    return mapping.get(exp_range, "")

# Helper: small normalizer for location (strip and return empty if blank)
def normalize_location(loc: str):
    if not loc:
        return ""
    return loc.strip()

# Helper: build a friendly query (keeps original query)
def build_job_query(title, location, experience_range, job_type):
    q = title.strip()
    if location:
        q += f" in {location.strip()}"
    if experience_range and experience_range != "Any":
        q += f" ({experience_range})"
    if job_type and job_type != "Any":
        q += f" [{job_type}]"
    return q

# Safe multi-page JSearch wrapper (returns list of dicts)
def search_jobs_rapidapi(query, location="", experience="", job_type="", posted="all", max_pages=10):
    KEY = os.getenv("RAPIDAPI_KEY", "")
    HOST = "jsearch.p.rapidapi.com"

    if not KEY:
        # If no key, return empty list (up to you to surface an error elsewhere)
        return []

    headers = {
        "x-rapidapi-key": KEY,
        "x-rapidapi-host": HOST
    }

    all_results = []
    for page in range(1, max_pages + 1):
        params = {
            "query": query,
            "page": page,
            "num_pages": 1,
            "date_posted": posted,
        }
        if location:
            params["location"] = location
        if job_type and job_type != "Any":
            params["employment_types"] = job_type.lower()
        if experience:
            params["experience_required"] = experience

        try:
            resp = requests.get(f"https://{HOST}/search", headers=headers, params=params, timeout=15)
            if resp.status_code != 200:
                break
            data = resp.json()
            chunk = data.get("data", []) or []
            # sanitize: only dicts
            chunk = [c for c in chunk if isinstance(c, dict)]
            if not chunk:
                break
            all_results.extend(chunk)
        except Exception:
            break

    return all_results

# Salary extractor (defensive)
def extract_salary(job):
    min_sal = job.get("job_min_salary")
    max_sal = job.get("job_max_salary")
    period = job.get("job_salary_period", "year")

    if min_sal and max_sal:
        try:
            return f"{int(min_sal):,} – {int(max_sal):,} / {period}"
        except:
            return f"{min_sal} – {max_sal} / {period}"

    est = job.get("estimated_salary")
    if est:
        return est

    sal = job.get("salary")
    if sal:
        return sal

    comp = job.get("compensation")
    if comp:
        return comp

    return "Not disclosed"

# Convert numeric amount (assumes amount is in USD if not specified). Uses env USD_TO_INR if present.
def convert_to_inr(amount):
    try:
        rate = float(os.getenv("USD_TO_INR", "83.0"))
        return int(round(float(amount) * rate))
    except:
        return int(amount)

# Salary benchmark: calls an optional RapidAPI salary endpoint (if SALARY_API_HOST present) else fallback
def get_salary_benchmark(job_title, city):
    # Make a best effort to call a RapidAPI salary API if configured
    host = os.getenv("SALARY_API_HOST", "").strip()  # e.g. "job-salary-data.p.rapidapi.com"
    key = os.getenv("RAPIDAPI_KEY", "")

    if host and key:
        headers = {
            "x-rapidapi-host": host,
            "x-rapidapi-key": key
        }
        # This is a generic best-effort path. Adapt to the specific API you're using.
        try:
            url = f"https://{host}/salary"
            params = {"jobTitle": job_title, "city": city}
            r = requests.get(url, headers=headers, params=params, timeout=12)
            if r.status_code == 200:
                data = r.json()
                # Example expected shape: {'salary_min': 50000, 'salary_max': 120000}
                if isinstance(data, dict) and ("salary_min" in data or "salaryMax" in data or "salary_max" in data):
                    # normalize keys
                    return {
                        "salary_min": data.get("salary_min") or data.get("salaryMin") or data.get("salary_min_est") or 0,
                        "salary_max": data.get("salary_max") or data.get("salaryMax") or data.get("salary_max_est") or 0,
                    }
        except Exception:
            pass

    # Fallback: return an AI-friendly sentinel so UI can use AI-based estimation
    return {"fallback": True}

# AI market-value stub (replace with your get_gemini_response or model call)
def ai_market_value(job_title, city, resume_text, salary_min=None, salary_max=None):
    # Replace this with an LLM call to produce a human-friendly evaluation
    # For now, make a simple heuristic
    try:
        avg = None
        if salary_min and salary_max:
            avg = (float(salary_min) + float(salary_max)) / 2
            resume_market_value = f"Estimated market value: ₹{int(avg):,}"
            expected_range = f"₹{int(salary_min):,} – ₹{int(salary_max):,}"
        else:
            resume_market_value = "Estimated market value: ₹8,00,000 – ₹12,00,000 (AI fallback)"
            expected_range = "₹8,00,000 – ₹12,00,000"

        negotiation_tip = "Your resume indicates strong skills; consider negotiating toward the top of the range if you can show impact metrics."

        return {
            "resume_market_value": resume_market_value,
            "expected_range": expected_range,
            "negotiation_tip": negotiation_tip
        }
    except Exception:
        return {"error": "AI evaluation failed"}


# --- Tesseract/API Config & Helper Functions ---
# --- Tesseract-OCR Path Configuration ---
TESSERACT_PATH = r'C:\Program Files\Tesseract-OCR\tesseract.exe' # Use the verified path
try:
    if os.path.exists(TESSERACT_PATH):
        pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
    else:
        pytesseract.pytesseract.tesseract_cmd = 'tesseract' # Fallback attempt
except Exception as e:
    pass 

# --- Environment Variable & Google AI Config ---
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

if not GOOGLE_API_KEY:
    pass
else:
    try:
        genai.configure(api_key=GOOGLE_API_KEY)
    except Exception as e:
        pass

# =========================================================================
# --- SESSION STATE INITIALIZATION (Must be early and UNCONDITIONAL) ---
# =========================================================================
if 'experience' not in st.session_state: st.session_state.experience = []
if 'education' not in st.session_state: st.session_state.education = []
if 'projects' not in st.session_state: st.session_state.projects = []
if 'achievements' not in st.session_state: st.session_state.achievements = []
if 'certificates' not in st.session_state: st.session_state.certificates = []
if 'custom_links' not in st.session_state: st.session_state.custom_links = []
if 'docx_buffer' not in st.session_state: st.session_state.docx_buffer = None
if 'pdf_buffer' not in st.session_state: st.session_state.pdf_buffer = None
if 'resume_filename' not in st.session_state: st.session_state.resume_filename = "resume"
if 'quality_response' not in st.session_state: st.session_state.quality_response = ""
if 'improved_resume_text' not in st.session_state: st.session_state.improved_resume_text = ""
if 'quality_filename' not in st.session_state: st.session_state.quality_filename = ""
if 'generated_cover_letter' not in st.session_state: st.session_state.generated_cover_letter = ""
if 'mock_interview_messages' not in st.session_state: st.session_state.mock_interview_messages = []
if 'mock_interview_jd' not in st.session_state: st.session_state.mock_interview_jd = ""
if 'mock_interview_error' not in st.session_state: st.session_state.mock_interview_error = None
if 'skill_gap_analysis_output' not in st.session_state: st.session_state.skill_gap_analysis_output = ""
if 'jd_analyzer_output' not in st.session_state: st.session_state.jd_analyzer_output = ""
if 'trend_jds' not in st.session_state: st.session_state.trend_jds = []
if 'trend_analysis_output' not in st.session_state: st.session_state.trend_analysis_output = ""
if 'predictive_qa_output' not in st.session_state: st.session_state.predictive_qa_output = ""
if 'portfolio_projects' not in st.session_state: st.session_state.portfolio_projects = []
if 'portfolio_generated_code' not in st.session_state: st.session_state.portfolio_generated_code = ""
if 'pf_name' not in st.session_state: st.session_state.pf_name = st.session_state.get('resume_name', '')
if 'pf_email' not in st.session_state: st.session_state.pf_email = st.session_state.get('resume_email', '')
if 'pf_linkedin' not in st.session_state: st.session_state.pf_linkedin = st.session_state.get('resume_linkedin', '')
if 'pf_github' not in st.session_state: st.session_state.pf_github = st.session_state.get('resume_github', '')
if 'pf_summary' not in st.session_state: st.session_state.pf_summary = st.session_state.get('resume_summary', '')
if 'pf_skills' not in st.session_state: st.session_state.pf_skills = st.session_state.get('resume_skills', '')
if 'pf_hero_text' not in st.session_state: st.session_state.pf_hero_text = ""
if 'active_view' not in st.session_state: st.session_state.active_view = "✍️ Resume Maker"
if 'resume_source_cl' not in st.session_state: st.session_state.resume_source_cl = "Upload a New Resume"
if 'resume_source_qa' not in st.session_state: st.session_state.resume_source_qa = "Upload a New Resume"
if 'api_response_tab1' not in st.session_state: st.session_state.api_response_tab1 = ""
if 'pdf_display_tab1' not in st.session_state: st.session_state.pdf_display_tab1 = ""
if 'resume_text_tab1' not in st.session_state: st.session_state.resume_text_tab1 = ""
if 'jd_text_tab1' not in st.session_state: st.session_state.jd_text_tab1 = ""
if 'uploaded_filename_tab1' not in st.session_state: st.session_state.uploaded_filename_tab1 = ""
if 'jd_multi' not in st.session_state: st.session_state.jd_multi = ""
if "chat_messages" not in st.session_state: st.session_state.chat_messages = []
if "job_results" not in st.session_state:
    st.session_state.job_results = []

if "job_page" not in st.session_state:
    st.session_state.job_page = 0


def generate_application_materials(job_title, job_description, company, resume_text):
    prompt = f"""
    You are an expert job application assistant.

    Generate the following for this job:

    JOB TITLE: {job_title}
    COMPANY: {company}
    JOB DESCRIPTION:
    {job_description}

    CANDIDATE RESUME:
    {resume_text}

    ---
    1. A tailored, highly personalized cover letter (150–220 words)
    2. Application answer set (4–6 Q&A):
       - Strengths
       - Weaknesses
       - Why should we hire you?
       - Explain a challenge you solved
       - Where do you see yourself in 5 years?
    3. A strong “Why do you want this job?” response (4–6 sentences)

    Make the tone confident, specific, achievement-based, and avoids generic clichés.
    """

    import google.generativeai as genai
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(prompt)

    return response.text

# =========================================================================
# --- CONSOLIDATED FUNCTION DEFINITIONS (Correct Order for Python) ---
# =========================================================================
def build_job_query(title, location, exp, job_type):
    q = title.strip()

    if location:
        q += f" in {location}"

    if exp and exp != "Any":
        q += f" {exp} experience"

    if job_type and job_type != "Any":
        q += f" {job_type.lower()}"

    return q.strip()


# --- Level 1: Core Helpers ---

def format_link_as_text(url, label=None):
    """Checks if a URL is valid and returns the URL string, otherwise returns None."""
    if not url: return None
    url = url.strip()
    
    # Check for mailto: or http(s):// prefix
    if url.lower().startswith("http://") or url.lower().startswith("https://") or url.lower().startswith("mailto:"):
        return url
    
    # Handle common cases where prefix is missing (e.g., just 'linkedin.com/in/...')
    if label and 'linkedin.com' in url.lower():
         return f"https://{url}" if not url.startswith('http') else url
    if label and '@' in url.lower(): # Handles email
         return f"mailto:{url}" if not url.startswith('mailto') else url
    if 'github.com' in url.lower(): # Handles GitHub
         return f"https://{url}" if not url.startswith('http') else url
         
    # If it seems like a URL but is missing http(s), assume https
    if '.' in url and len(url) > 5:
        return f"https://{url}"
        
    return None # Return None if not a valid/detectable link

def get_gemini_response(input_prompt, temperature=0.0):
    """ Calls the Gemini API for analysis/ranking (deterministic). """
    if not GOOGLE_API_KEY: return "Error: API Key not configured."
    try:
        model = genai.GenerativeModel('models/gemini-pro-latest')
        config = GenerationConfig(temperature=temperature)
        response = model.generate_content(input_prompt, generation_config=config)
        
        # --- FIX: Check for empty parts even if the finish_reason is STOP (1) ---
        if not response.parts:
            # Check if the block was due to safety or prompt issues
            if response.prompt_feedback and response.prompt_feedback.block_reason:
                return f"Error: Request blocked by API due to {response.prompt_feedback.block_reason.name}."
            # Check for general empty response (The error scenario)
            if response.candidates and response.candidates[0].finish_reason != 1:
                return f"Error: Response generation stopped due to {response.candidates[0].finish_reason.name}."
            
            # If the response finished (reason 1) but has no parts, treat it as a temporary block
            return "Error: Empty response received from API (possible content safety block)."
        # ----------------------------------------------------------------------
        
        # Original logic check
        if response.candidates and response.candidates[0].finish_reason != 1: # 1 = STOP
            return f"Error: Response generation stopped due to {response.candidates[0].finish_reason.name}."

        return response.text
    except Exception as e:
        error_message = f"Error calling Gemini: {e}"
        if "API key not valid" in str(e): error_message = "Error: Google API key is not valid. Please check your .env file."
        elif "quota" in str(e).lower():
            delay_match = re.search(r'retry_delay {\s*seconds:\s*(\d+)\s*}', str(e))
            delay_msg = f" Please wait {delay_match.group(1)} seconds and try again." if delay_match else " Try again shortly."
            error_message = f"Error: Rate limit exceeded (Free Tier?).{delay_msg}"
        elif "candidate" in str(e).lower() and "finish_reason: SAFETY" in str(e):
            error_message = "Error: The response was blocked due to safety concerns."
        return error_message


def get_gemini_response_chat(prompt_list, temperature=0.7):
    """ Calls the Gemini API for chat (creative, streaming). """
    if not GOOGLE_API_KEY: return iter(["Error: API Key not configured."])
    try:
        model = genai.GenerativeModel('gemini-pro-latest')
        config = GenerationConfig(temperature=temperature)
        response = model.generate_content(prompt_list, generation_config=config, stream=True)
        return response
    except Exception as e:
        error_message = f"Error calling Gemini Chat: {e}"
        if "API key not valid" in str(e): error_message = "Error: Google API key is not valid."
        elif "quota" in str(e).lower(): error_message = "Error: Rate limit exceeded for chat."
        elif "candidate" in str(e).lower() and "finish_reason: SAFETY" in str(e):
            error_message = "Error: Chat response blocked due to safety concerns."
        st.error(error_message)
        return iter([f"Sorry, an error occurred: {error_message}"])


def extract_text_from_file(uploaded_file):
    """ Extracts text from PDF (text/scanned) and DOCX. """
    file_bytes = uploaded_file.getvalue()
    file_name = uploaded_file.name
    text = ""

    try:
        if file_name.endswith('.pdf'):
            try:
                pdf_reader = pdf.PdfReader(io.BytesIO(file_bytes))
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text: text += page_text + "\n"
            except Exception as text_extract_err:
                 text = ""

            if not text or len(text.strip()) < 100:
                text = ""
                try:
                    images = convert_from_bytes(file_bytes, fmt='png', thread_count=4)
                    extracted_pages = 0
                    for i, image in enumerate(images):
                        try:
                            page_ocr_text = pytesseract.image_to_string(image)
                            if page_ocr_text:
                                text += page_ocr_text + "\n"
                                extracted_pages += 1
                        except pytesseract.TesseractNotFoundError:
                            return None
                        except Exception as ocr_page_err:
                            pass
                except Exception as ocr_err:
                     return None

        elif file_name.endswith('.docx'):
            doc = docx.Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs: text += para.text + "\n"

        return text.strip() if text else None

    except Exception as e:
        st.error(f"Error processing file '{file_name}': {e}", icon="🚨")
        return None

def clean_json_response(text):
    """ Extracts the first valid JSON object or list from text, handling markdown. """
    match_start = re.search(r'[{[]', text)
    if not match_start:
        raise json.JSONDecodeError("No JSON object or array start found.", text, 0)
    start_index = match_start.start()
    start_char = text[start_index]
    end_char = '}' if start_char == '{' else ']'
    open_count = 0
    in_string = False
    escaped = False
    end_index = -1
    for i in range(start_index, len(text)):
        char = text[i]
        if in_string:
            if char == '"' and not escaped: in_string = False
            escaped = (char == '\\' and not escaped)
        else:
            if char == '"': in_string = True; escaped = False
            elif char == start_char: open_count += 1
            elif char == end_char:
                open_count -= 1
                if open_count == 0: end_index = i + 1; break
        if char != '\\': escaped = False
    if end_index == -1: raise json.JSONDecodeError(f"Mismatched JSON braces.", text, len(text))
    clean_text = text[start_index:end_index]
    try:
        json.loads(clean_text); return clean_text
    except json.JSONDecodeError as e:
       raise json.JSONDecodeError(f"Failed to parse cleaned JSON: {e.msg}", clean_text, e.pos)

def build_annotated_text(text, keywords):
    """ Builds list for st_annotated_text highlighting keywords. """
    if not text or not keywords: return [text or ""]
    annotated_list = []
    parts = re.split(r'(\W+)', text) # Split on non-word characters
    parts = [p for p in parts if p]
    keywords_lower = {k.lower() for k in keywords if isinstance(k, str)}
    for part in parts:
        part_lower = part.lower()
        if part_lower in keywords_lower:
            annotated_list.append((part, "MATCH", "#59AA7B"))
        else:
            annotated_list.append(part)
    # Combine adjacent non-match strings
    combined_list = []; current_string = ""
    for item in annotated_list:
        if isinstance(item, tuple):
            if current_string: combined_list.append(current_string)
            combined_list.append(item); current_string = ""
        else: current_string += item
    if current_string: combined_list.append(current_string)
    return combined_list

def clean_html_response(text):
    """ Cleans the AI's response to get only the raw HTML code. """
    match = re.search(r'```html?(.*)```', text, re.DOTALL | re.IGNORECASE)
    if match: return match.group(1).strip()
    if text.strip().startswith('<') and text.strip().endswith('>'): return text.strip()
    st.warning("AI response did not contain clear HTML code.", icon="⚠️")
    return "" # Return empty string if no clear HTML

def format_resume_data_for_prompt(data):
    """Converts the resume data dictionary into a clean text format for the LLM."""
    try:
        text = f"Candidate Name: {data.get('name', 'N/A')}\n"
        text += f"Location: {data.get('location', 'N/A')}\n"
        text += f"Contact: {data.get('phone', 'N/A')} | {data.get('email', 'N/A')}\n"
        links = [data.get('linkedin'), data.get('github')]
        links.extend([f"{link.get('label', 'Link')}: {link.get('url', '')}" for link in data.get('custom_links', [])])
        text += "Links: " + " | ".join(filter(None, links)) + "\n\n"
        if data.get('summary'): text += f"--- PROFESSIONAL SUMMARY ---\n{data['summary']}\n\n"
        if data.get('skills'): text += f"--- SKILLS ---\n{data['skills']}\n\n"
        if data.get('experience'):
            text += "--- EXPERIENCE ---\n"
            for exp in data['experience']:
                text += f"\nJob Title: {exp.get('title', 'N/A')}\n"
                text += f"Company: {exp.get('company', 'N/A')} | Dates: {exp.get('dates', 'N/A')}\n"
                if exp.get('description'):
                    text += "Accomplishments:\n"
                    for line in exp['description'].split('\n'):
                        line_strip = line.strip().lstrip('*- ');
                        if line_strip: text += f"- {line_strip}\n"
        if data.get('education'):
            text += "\n--- EDUCATION ---\n"
            for edu in data['education']:
                text += f"\nDegree: {edu.get('degree', 'N/A')}\n"
                text += f"Institution: {edu.get('institution', 'N/A')} | Dates: {edu.get('dates', 'N/A')}\n"
                if edu.get('details'): text += f"Details: {edu['details']}\n"
        if data.get('projects'):
            text += "\n--- PROJECTS ---\n"
            for proj in data['projects']:
                text += f"\nProject: {proj.get('name', 'N/A')}\n"
                if proj.get('description'):
                    for line in proj['description'].split('\n'):
                        line_strip = line.strip().lstrip('*- ');
                        if line_strip: text += f"- {line_strip}\n"
        if data.get('achievements'):
            text += "\n--- ACHIEVEMENTS & CERTIFICATES ---\n"
            for ach in data['achievements']:
                if ach.get('name'): text += f"- {ach['name']}\n"
        return text.strip()
    except Exception as e:
        st.error(f"Error formatting resume data for prompt: {e}", icon="🚨"); return None
       
        # --- The Function Definition to Paste ---
def create_analysis_report_docx(data, filename):
    """Generates a detailed, structured DOCX analysis report."""
    try:
        doc = docx.Document()
        doc.add_heading('ATS Analysis Report', level=0)
        doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Candidate File: {filename}")
        doc.add_paragraph()

        # Core Metrics
        doc.add_heading('1. Core Metrics', level=1)
        doc.add_paragraph(f"ATS Match Score: {data.get('JD_Match', 'N/A')}")
        doc.add_paragraph(f"Profile Summary: {data.get('Profile_Summary', 'N/A')}")
        
        # Keywords
        doc.add_heading('2. Keyword Match', level=1)
        doc.add_heading('Matched Keywords', level=2)
        doc.add_paragraph(", ".join(data.get('Matched_Keywords', ['None found.'])))
        doc.add_heading('Missing Keywords', level=2)
        doc.add_paragraph(", ".join(data.get('MissingKeywords', ['None found.'])))

        # Recruiter Flags
        doc.add_heading('3. Recruiter Flags', level=1)
        doc.add_heading('Green Flags (Positives)', level=2)
        for flag in data.get('Green_Flags', ['None identified.']):
            doc.add_paragraph(flag, style='List Bullet')
        doc.add_heading('Red Flags (Concerns)', level=2)
        for flag in data.get('Red_Flags', ['None identified.']):
            doc.add_paragraph(flag, style='List Bullet')

        # Suggested Interview Questions
        doc.add_heading('4. Suggested Interview Questions', level=1)
        doc.add_heading('Technical', level=2)
        for q in data.get('Interview_Questions', {}).get('Technical', ['None suggested.']):
            doc.add_paragraph(q, style='List Bullet')
        doc.add_heading('Behavioral', level=2)
        for q in data.get('Interview_Questions', {}).get('Behavioral', ['None suggested.']):
            doc.add_paragraph(q, style='List Bullet')

        # Suggested Summary
        doc.add_heading('5. Suggested Resume Summary', level=1)
        doc.add_paragraph(data.get('Suggested_Resume_Summary', 'N/A'))

        buffer = io.BytesIO(); doc.save(buffer); buffer.seek(0); return buffer
    except Exception as e:
        st.error(f"Error generating analysis report DOCX: {e}", icon="🚨"); return None

# =========================================================================
# --- LEVEL 2: MARKDOWN/COVER LETTER HELPERS (Define before Level 3) ---
# =========================================================================

def create_markdown_docx(markdown_text, title="AI Generated Report"):
    """Generates a simple DOCX file from a markdown string (used for CL/JD/QA)."""
    try:
        doc = Document()
        doc.add_heading(title, level=0)
        
        lines = markdown_text.split('\n')
        for line in lines:
            if line.startswith('## '):
                doc.add_heading(line.lstrip('# ').strip(), level=2)
            elif line.startswith('### '):
                doc.add_heading(line.lstrip('# ').strip(), level=3)
            elif line.startswith('* '):
                doc.add_paragraph(line.lstrip('* ').strip(), style='List Bullet')
            elif line.startswith('- '):
                doc.add_paragraph(line.lstrip('- ').strip(), style='List Bullet')
            elif line.strip():
                doc.add_paragraph(line.strip())
        
        buffer = io.BytesIO(); doc.save(buffer); buffer.seek(0); return buffer
    except Exception as e:
        st.error(f"Error generating markdown DOCX: {e}", icon="🚨"); return None

def create_markdown_pdf(markdown_text, title="AI Generated Report"):
    """Generates a professional PDF report from a markdown string, resolving all style conflicts."""
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
        story = []; styles = getSampleStyleSheet()
        
        # Define custom styles using a unique prefix ('ATS') to guarantee no conflict.
        ATS_BODY = ParagraphStyle(name='ATS_Body_Text', parent=styles['Normal'], spaceAfter=6, leading=14, fontSize=10)
        styles.add(ATS_BODY)

        styles.add(ParagraphStyle(name='ATS_Heading0', parent=styles['Normal'], fontSize=16, spaceBefore=12, spaceAfter=6, fontName='Helvetica-Bold'))
        styles.add(ParagraphStyle(name='ATS_Heading2', parent=styles['Normal'], fontSize=12, spaceBefore=10, spaceAfter=4, fontName='Helvetica-Bold', textColor='#003366'))
        styles.add(ParagraphStyle(name='ATS_Heading3', parent=styles['Normal'], fontSize=11, spaceBefore=8, spaceAfter=2, fontName='Helvetica-Bold'))
        styles.add(ParagraphStyle(name='ATS_BulletPoint', parent=ATS_BODY, leftIndent=0.25*inch, firstLineIndent=0, bulletIndent=0.15*inch, bulletText='•'))

        # Add the main title
        story.append(Paragraph(title, styles['ATS_Heading0']))

        # Process the markdown text
        lines = markdown_text.split('\n')
        for line in lines:
            line_strip = line.strip()
            if not line_strip:
                continue
            elif line_strip.startswith('## '):
                story.append(Paragraph(line_strip.lstrip('## '), styles['ATS_Heading2']))
            elif line_strip.startswith('### '):
                story.append(Paragraph(line_strip.lstrip('### '), styles['ATS_Heading3']))
            elif line_strip.startswith('* ') or line_strip.startswith('- '):
                story.append(Paragraph(line_strip.lstrip('*- '), styles['ATS_BulletPoint']))
            else:
                story.append(Paragraph(line_strip, styles['ATS_Body_Text']))
        
        doc.build(story); buffer.seek(0); return buffer
    except Exception as e:
        st.error(f"Error generating PDF from markdown: {e}", icon="🚨"); 
        raise RuntimeError(f'ReportLab Final Error: {e}') from e

# -------------------------------------------------------------------------
# --- LEVEL 3: SPECIALIZED DOCUMENT BUILDERS (Need Level 1 defined first) ---
# -------------------------------------------------------------------------

def create_cover_letter_docx(letter_text):
    """Generates a simple .docx file in memory from cover letter text."""
    # Calls create_markdown_docx
    return create_markdown_docx(letter_text, title="Cover Letter")

def create_cover_letter_pdf(letter_text):
    """Generates a simple .pdf file in memory from cover letter text using ReportLab."""
    try:
        buffer = io.BytesIO(); doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
        story = []; styles = getSampleStyleSheet(); body_style = ParagraphStyle(name='BodyStyle', parent=styles['Normal'], spaceAfter=6, leading=16)
        formatted_text = letter_text.replace('\n\n', '<br/><br/>').replace('\n', ' ')
        story.append(Paragraph(formatted_text, body_style))
        doc.build(story); buffer.seek(0); return buffer
    except Exception as e:
        st.error(f"Error generating cover letter .pdf: {e}", icon="🚨"); return None

# --- Level 3: RESUME BUILDERS (These must be defined after format_link_as_text and other helpers) ---
def create_resume_docx(data):
    """Generates an ATS-friendly .docx file in memory with clickable links and styled headings."""
    try:
        doc = Document()
        for section in doc.sections:
            # Set minimal margins for maximum content space
            section.left_margin = Inches(0.5)    # Reduced from 1.0 inch
            section.right_margin = Inches(0.5)   # Reduced from 1.0 inch
            section.top_margin = Inches(0.25)    # Reduced from 0.5 inch
            section.bottom_margin = Inches(0.5)  # Reduced from 1.0 inch
            section.header_distance = Inches(0)  # Remove header space
            section.footer_distance = Inches(0)  # Remove footer space
            
            # Remove headers and footers
            for header in section.header.paragraphs:
                header.clear()
            for footer in section.footer.paragraphs:
                footer.clear()
            section.different_first_page_header_footer = False
        # Set default font to Times New Roman
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(10)  # Slightly smaller than Calibri for better readability

        # Helper function to insert hyperlinked text (required for clickable links)
        def add_hyperlink(paragraph, text, url):
            part = paragraph.part
            # Ensure proper URL structure for external link definition
            url = url if url.lower().startswith(('http', 'mailto')) else 'http://' + url 
            rId = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
            hyperlink.set(docx.oxml.ns.qn('r:id'), rId, )
            new_run = docx.oxml.shared.OxmlElement('w:r')
            rPr = docx.oxml.shared.OxmlElement('w:rPr')
            new_run.append(rPr)
            new_run.text = text
            hyperlink.append(new_run)
            paragraph._element.append(hyperlink)
            return hyperlink
            
        # Helper to create a styled heading block (FIXED XML PARSING)
        def add_styled_heading_docx(document, text):
            # Add some space before the heading for visual separation
            document.add_paragraph(style='Normal').add_run().add_break() # Adds a blank line

            table = document.add_table(rows=1, cols=1)
            table.width = Inches(6.5) 
            cell = table.cell(0, 0)
            
            # --- FIX: Simplified XML Shading Element ---
            # Define shading directly using the hex color D9D9D9 (light grey)
            shading_elm = docx.oxml.shared.OxmlElement('w:shd')
            shading_elm.set(docx.oxml.ns.qn('w:val'), 'clear')
            shading_elm.set(docx.oxml.ns.qn('w:fill'), 'D9D9D9')
            cell._tc.get_or_add_tcPr().append(shading_elm)
            # --- END FIX ---

            # Add text
            p = cell.paragraphs[0]
            for run in list(p.runs): run._element.getparent().remove(run._element)
            
            run = p.add_run(text.upper())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:ascii'), 'Times New Roman')
            run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:hAnsi'), 'Times New Roman')
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Minimize vertical padding for compact layout
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            # Set heading styles to Times New Roman with compact spacing
            for i in range(1, 4):
                style_name = f'Heading {i}'
                if style_name in doc.styles:
                    style = doc.styles[style_name]
                    style.font.name = 'Times New Roman'
                    style._element.rPr.rFonts.set(docx.oxml.ns.qn('w:ascii'), 'Times New Roman')
                    style._element.rPr.rFonts.set(docx.oxml.ns.qn('w:hAnsi'), 'Times New Roman')
                    style.paragraph_format.space_before = Pt(4)  # Reduced from default
                    style.paragraph_format.space_after = Pt(2)   # Reduced from default
                    if i == 1:
                        style.font.size = Pt(12)  # Slightly smaller for compactness
                    else:
                        style.font.size = Pt(11 - (i-1))  # Smaller progressioncal padding for compact layout

        # Name and Contact Header
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(data.get('name', 'NAME MISSING').upper())
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:ascii'), 'Times New Roman')
        run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:hAnsi'), 'Times New Roman')
        run.font.size = Pt(18)
        run.font.bold = True
        
        p_contact = doc.add_paragraph(); p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_contact.paragraph_format.space_after = Pt(0)
        
        contact_line_parts = []
        if data.get('phone'): contact_line_parts.append(data.get('phone'))
        if data.get('location'): contact_line_parts.append(data.get('location'))
        
        email_display_text = data.get('email')
        if email_display_text:
            email_valid_url = format_link_as_text(email_display_text, label='email')
            if email_valid_url:
                add_hyperlink(p_contact, email_display_text, email_valid_url)
            else:
                contact_line_parts.append(email_display_text)

        if contact_line_parts:
            p_contact.add_run(" | ".join(contact_line_parts)).font.size = Pt(10)


        # Add Links (LinkedIn, GitHub, Custom) in a new paragraph
        p_links = doc.add_paragraph(); p_links.alignment = WD_ALIGN_PARAGRAPH.CENTER
        link_elements = []

        for label, url in [('LinkedIn', data.get('linkedin')), ('GitHub', data.get('github'))]:
            valid_url = format_link_as_text(url, label=label)
            if valid_url:
                display_text = label
                if label=='GitHub' and url and url.split('/')[-1]:
                    display_text = f"GitHub ({url.split('/')[-1]})"
                link_elements.append((display_text, valid_url))

        for link_obj in data.get('custom_links', []):
            if link_obj.get('label') and link_obj.get('url'):
                valid_url = format_link_as_text(link_obj.get('url'))
                if valid_url:
                    link_elements.append((link_obj['label'], valid_url))
        
        for i, (text, url) in enumerate(link_elements):
            add_hyperlink(p_links, text, url)
            if i < len(link_elements) - 1:
                p_links.add_run(" | ")

        # --- Main Sections with Styled Headings ---

        # 1. Professional Summary
        if data.get('summary'): 
            add_styled_heading_docx(doc, 'PROFESSIONAL SUMMARY')
            doc.add_paragraph(data['summary'])
            doc.add_paragraph()  # Add space after section
        
        # 2. Skills
        if data.get('skills'): 
            add_styled_heading_docx(doc, 'SKILLS')
            doc.add_paragraph(data['skills'])
            doc.add_paragraph()  # Add space after section
            
        # 5. Education (moved up to position 3)
        if data.get('education'):
            add_styled_heading_docx(doc, 'EDUCATION')
            for edu in data['education']:
                doc.add_heading(edu.get('degree', 'N/A'), level=2)
                p_edu = doc.add_paragraph()
                p_edu.add_run(edu.get('institution', 'N/A')).bold = True
                p_edu.add_run(f"\t | \t{edu.get('dates', 'N/A')}")
                if edu.get('details'): 
                    doc.add_paragraph(edu['details'])
            doc.add_paragraph()  # Add space after section
        
        # 3. Experience (moved to position 4)
        if data.get('experience'):
            add_styled_heading_docx(doc, 'EXPERIENCE')
            for exp in data['experience']:
                doc.add_heading(exp.get('title', 'N/A'), level=2)
                p_company = doc.add_paragraph()
                p_company.add_run(exp.get('company', 'N/A')).bold = True
                p_company.add_run(f"\t | \t{exp.get('dates', 'N/A')}")
                if exp.get('description'):
                    for bullet in exp['description'].split('\n'):
                        bullet_text = bullet.strip().lstrip('*- ')
                        if bullet_text: 
                            doc.add_paragraph(bullet_text, style='List Bullet')
            doc.add_paragraph()  # Add space after section
            
        # 6. Projects (position 5)
        if data.get('projects'):
            add_styled_heading_docx(doc, 'PROJECTS')
            for proj in data['projects']:
                doc.add_heading(proj.get('name', 'N/A'), level=2)
                if proj.get('description'):
                    for bullet in proj['description'].split('\n'):
                        bullet_text = bullet.strip().lstrip('*- ')
                        if bullet_text:
                            if bullet_text.startswith("http"): 
                                add_hyperlink(doc.add_paragraph(), bullet_text, bullet_text)
                            else: 
                                doc.add_paragraph(bullet_text, style='List Bullet')
            doc.add_paragraph()  # Add space after section
        
        if data.get('achievements'):
            add_styled_heading_docx(doc, 'ACHIEVEMENTS')
            for item in data['achievements']:
                if item.get('name'): doc.add_paragraph(item['name'], style='List Bullet')
                
        if data.get('certificates'):
            add_styled_heading_docx(doc, 'CERTIFICATES')
            for cert in data['certificates']:
                if cert.get('name'):
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(cert['name']).bold = True
                    if cert.get('issuer') or cert.get('date'):
                        details = []
                        if cert.get('issuer'):
                            details.append(cert['issuer'])
                        if cert.get('date'):
                            details.append(cert['date'])
                        if details:
                            p.add_run('\n' + ' | '.join(details)).font.size = Pt(9)
                            p.paragraph_format.space_after = Pt(4)
        
        buffer = io.BytesIO(); doc.save(buffer); buffer.seek(0); return buffer
    except Exception as e:
        st.error(f"Error generating .docx file: {e}", icon="🚨"); return None

def create_resume_pdf(data):
    """Generates an ATS-friendly .pdf file in memory with clickable links and styled headings (ReportLab)."""
    try:
        buffer = io.BytesIO()
        # FIX: Reduced top margin from 0.75 to 0.5 inches
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=0.75*inch, leftMargin=0.75*inch, topMargin=0.5*inch, bottomMargin=0.75*inch)
        story = []; styles = getSampleStyleSheet()
        
        # Custom Styles (using unique RESUME_ prefix) - Using Times New Roman
        styles.add(ParagraphStyle(
            name='RESUME_Name', 
            parent=styles['h1'], 
            alignment=TA_CENTER, 
            fontSize=18, 
            spaceAfter=4, 
            fontName='Times-Bold'  # Changed to Times-Bold
        ))
        
        styles.add(ParagraphStyle(
            name='RESUME_Contact', 
            parent=styles['Normal'], 
            alignment=TA_CENTER, 
            fontSize=9, 
            spaceAfter=2,
            fontName='Times-Roman'  # Changed to Times-Roman
        ))
        
        styles.add(ParagraphStyle(
            name='RESUME_SubHeading', 
            parent=styles['Normal'], 
            fontSize=11, 
            spaceBefore=6, 
            spaceAfter=2, 
            fontName='Times-Bold'  # Changed to Times-Bold
        ))
        
        # Updated body style with full justification and Times New Roman
        RESUME_Body = ParagraphStyle(
            name='RESUME_Body', 
            parent=styles['Normal'], 
            fontSize=10, 
            spaceAfter=6, 
            leading=14,
            alignment=TA_JUSTIFY,  # Full justification
            wordWrap='LTR',
            splitLongWords=True,
            spaceShrinkage=0.05,
            trailingSpace=1,
            fontName='Times-Roman'  # Changed to Times-Roman
        )
        styles.add(RESUME_Body)
        
        # Updated bullet style with proper indentation and alignment
        styles.add(ParagraphStyle(
            name='RESUME_Bullet', 
            parent=RESUME_Body, 
            leftIndent=0.25*inch, 
            firstLineIndent=-0.15*inch,  # Negative indent for bullet
            spaceBefore=0, 
            spaceAfter=2, 
            bulletIndent=0.15*inch, 
            bulletText='•',
            alignment=TA_LEFT,  # Left align bullet points
            fontName='Times-Roman'  # Changed to Times-Roman
        ))
        
        # Job details style with proper spacing
        styles.add(ParagraphStyle(
            name='RESUME_JobDetails', 
            parent=styles['Normal'], 
            fontSize=10, 
            spaceAfter=2,
            alignment=TA_LEFT,
            fontName='Times-Roman'  # Changed to Times-Roman
        ))
        
        # Helper to create a styled heading block for PDF (like the image)
        def add_styled_heading_pdf(text_content):
            from reportlab.lib import colors
            from reportlab.platypus import Table, TableStyle
            
            # Use HexColor for the specific grey (D9D9D9)
            grey_fill = HexColor('#D9D9D9')

            heading_para = Paragraph(text_content.upper(), 
                                     ParagraphStyle(name='HeadingStyle', parent=styles['Normal'], 
                                                    fontSize=12, fontName='Helvetica-Bold', 
                                                    textColor=colors.black, alignment=TA_LEFT,
                                                    spaceBefore=4, spaceAfter=4))
            
            # Create a one-cell table for the background color
            table_data = [[heading_para]]
            table = Table(table_data, colWidths=[6.5*inch]) # Match docx width
            
            table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,-1), grey_fill),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('LEFTPADDING', (0,0), (-1,-1), 8), # Add some internal padding
                ('RIGHTPADDING', (0,0), (-1,-1), 8),
                ('TOPPADDING', (0,0), (-1,-1), 4),
                ('BOTTOMPADDING', (0,0), (-1,-1), 4),
            ]))
            return [Spacer(1, 0.05*inch), table, Spacer(1, 0.08*inch)] # FIX: Reduced top space before the box


        # Name Header 
        story.append(Paragraph(data.get('name', 'NAME MISSING').upper(), styles['RESUME_Name']))

        # Contact/Link Block 
        contact_line_parts = []
        
        # Location/Phone (non-link text)
        if data.get('location'): contact_line_parts.append(data.get('location'))
        if data.get('phone'): contact_line_parts.append(data.get('phone'))

        # Email Link
        email_url = format_link_as_text(data.get('email'), label='email')
        if email_url: contact_line_parts.append(f"<a href='{email_url}'>{data.get('email')}</a>")
        elif data.get('email'): contact_line_parts.append(data.get('email'))

        if contact_line_parts: story.append(Paragraph(" | ".join(contact_line_parts), styles['RESUME_Contact']))

        # Social/Custom Links
        link_line_parts = []
        for label, url in [('LinkedIn', data.get('linkedin')), ('GitHub', data.get('github'))]:
            valid_url = format_link_as_text(url, label=label)
            if valid_url:
                link_line_parts.append(f"<a href='{valid_url}'>{label}</a>")

        for link_obj in data.get('custom_links', []):
            if link_obj.get('label') and link_obj.get('url'):
                valid_url = format_link_as_text(link_obj.get('url'))
                if valid_url:
                    link_line_parts.append(f"<a href='{valid_url}'>{link_obj['label']}</a>")

        if link_line_parts: story.append(Paragraph(" | ".join(link_line_parts), styles['RESUME_Contact']))
        
        story.append(Spacer(1, 0.05*inch)) # Reduced vertical space after links
        
        # --- Main Sections with Styled Headings ---
        
        # 1. Professional Summary
        if data.get('summary'):
            story.extend(add_styled_heading_pdf('PROFESSIONAL SUMMARY'))
            story.append(Paragraph(data['summary'].replace('\n', '<br/>'), styles['RESUME_Body']))
            story.append(Spacer(1, 0.08*inch))
        
        # 2. Skills
        if data.get('skills'):
            story.extend(add_styled_heading_pdf('SKILLS'))
            story.append(Paragraph(data['skills'], styles['RESUME_Body']))
            story.append(Spacer(1, 0.08*inch))
            
        # 5. Education (moved up to position 3)
        if data.get('education'):
            story.extend(add_styled_heading_pdf('EDUCATION'))
            for edu in data['education']:
                story.append(Paragraph(edu.get('degree', 'N/A'), styles['RESUME_SubHeading']))
                story.append(Paragraph(f"<b>{edu.get('institution', 'N/A')}</b> | {edu.get('dates', 'N/A')}", styles['RESUME_JobDetails']))
                if edu.get('details'):
                    details_lines = [line.strip().lstrip('*- ') for line in edu['details'].split('\n') if line.strip()]
                    for line in details_lines: story.append(Paragraph(line, styles['RESUME_Bullet']))
            story.append(Spacer(1, 0.08*inch))
        
        # 3. Experience (moved to position 4)
        if data.get('experience'):
            story.extend(add_styled_heading_pdf('EXPERIENCE'))
            for exp in data['experience']:
                story.append(Paragraph(exp.get('title', 'N/A'), styles['RESUME_SubHeading']))
                story.append(Paragraph(f"<b>{exp.get('company', 'N/A')}</b> | {exp.get('dates', 'N/A')}", styles['RESUME_JobDetails']))
                desc_text = exp.get('description', '')
                if desc_text:
                    lines = [line.strip().lstrip('*- ') for line in desc_text.split('\n') if line.strip()]
                    for line in lines: story.append(Paragraph(line, styles['RESUME_Bullet']))
                story.append(Spacer(1, 0.08*inch))
            
        # 6. Projects (position 5)
        if data.get('projects'):
            story.extend(add_styled_heading_pdf('PROJECTS'))
            for proj in data['projects']:
                story.append(Paragraph(proj.get('name', 'N/A'), styles['RESUME_SubHeading']))
                desc_text = proj.get('description', '')
                if desc_text:
                    lines = [line.strip().lstrip('*- ') for line in desc_text.split('\n') if line.strip()]
                    for line in lines:
                        if line.startswith("http"): 
                            valid_url = format_link_as_text(line)
                            story.append(Paragraph(f"<a href='{valid_url}'>{line}</a>", styles['RESUME_Body']))
                        else: 
                            story.append(Paragraph(line, styles['RESUME_Bullet']))
                story.append(Spacer(1, 0.08*inch))
                
        if data.get('achievements'):
            story.extend(add_styled_heading_pdf('ACHIEVEMENTS'))
            for item in data['achievements']:
                if item.get('name'): story.append(Paragraph(item['name'], styles['RESUME_Bullet']))
                
        if data.get('certificates'):
            story.extend(add_styled_heading_pdf('CERTIFICATES'))
            for cert in data['certificates']:
                if cert.get('name'):
                    cert_text = f"<b>{cert['name']}</b>"
                    if cert.get('issuer') or cert.get('date'):
                        details = []
                        if cert.get('issuer'):
                            details.append(f"{cert['issuer']}")
                        if cert.get('date'):
                            details.append(f"{cert['date']}")
                        cert_text += f"<br/><font size='9' color='#666666'>{' | '.join(details)}</font>"
                    story.append(Paragraph(cert_text, styles['RESUME_Body']))
                    story.append(Spacer(1, 0.08*inch))
        
        try:
            doc.build(story); buffer.seek(0); return buffer
        except Exception as build_err:
             st.error(f"Error building PDF: {build_err}", icon="🚨"); print(f"!!! PDF Build Error: {build_err}"); return None
    except Exception as e:
        st.error(f"General error in create_resume_pdf: {e}", icon="🚨"); print(f"!!! General PDF Error: {e}"); return None

def create_cover_letter_docx(letter_text):
    """Generates a simple .docx file in memory from cover letter text."""
    return create_markdown_docx(letter_text, title="Cover Letter")

def create_cover_letter_pdf(letter_text):
    """Generates a simple .pdf file in memory from cover letter text using ReportLab."""
    try:
        buffer = io.BytesIO(); doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
        story = []; styles = getSampleStyleSheet(); body_style = ParagraphStyle(name='BodyStyle', parent=styles['Normal'], spaceAfter=6, leading=16)
        formatted_text = letter_text.replace('\n\n', '<br/><br/>').replace('\n', ' ')
        story.append(Paragraph(formatted_text, body_style))
        doc.build(story); buffer.seek(0); return buffer
    except Exception as e:
        st.error(f"Error generating cover letter .pdf: {e}", icon="🚨"); return None

def create_ranking_pdf(ranking_df, job_title="Candidate Ranking Report"):
    """Generates a structured PDF report from the ranking DataFrame with improved layout and justification."""
    try:
        buffer = io.BytesIO()
        
        # Use landscape orientation for better table display
        from reportlab.lib.pagesizes import landscape, A4
        doc = SimpleDocTemplate(
            buffer,
            pagesize=landscape(A4),
            rightMargin=30,
            leftMargin=30,
            topMargin=30,
            bottomMargin=30
        )
        
        story = []
        styles = getSampleStyleSheet()
        
        # Title with better styling
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=20,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#2c3e50')
        )
        story.append(Paragraph(f"<b>Ranking Report:</b> {job_title}", title_style))
        
        # Table data with wrapped text in Paragraphs
        header = ["Rank", "Candidate Filename", "ATS Score", "Confidence", "Missing Skills", "Recruiter Justification"]
        data = [header]

        # Define styles for table cells with proper justification
        cell_style = ParagraphStyle(
            'Cell',
            parent=styles['Normal'],
            fontSize=9,
            leading=11,
            spaceAfter=6,
            wordWrap='CJK',
            hyphenationLang='en-US',
            alignment=TA_JUSTIFY
        )
        
        # Add data rows with proper formatting
        for _, row in ranking_df.iterrows():
            # Clean the skills list for the cell
            skills = ", ".join(row.get('Missing_Critical_Skills', [])) if isinstance(row.get('Missing_Critical_Skills', []), list) else str(row.get('Missing_Critical_Skills', ''))
            
            # Wrap all cell content in Paragraphs for proper text wrapping and justification
            data.append([
                str(row.get('Rank', 'N/A')),
                Paragraph(str(row.get('Filename', 'N/A')), cell_style),
                str(row.get('Score', 'N/A')),
                str(row.get('Confidence', 'N/A')),
                Paragraph(skills, cell_style),
                Paragraph(str(row.get('Justification', 'N/A')), cell_style)
            ])
        
        # Create table with optimized column widths (total width ~740 for landscape A4)
        col_widths = [40, 130, 60, 100, 180, 250]
        table = Table(data, colWidths=col_widths, repeatRows=1)
        
        # Apply table styles
        table.setStyle(TableStyle([
            # Header row
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#003366')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            
            # Data rows
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor('#2c3e50')),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#dddddd')),
            ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#dddddd')),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        
        # Add alternating row colors for better readability
        for i in range(1, len(data)):
            if i % 2 == 0:  # Even rows
                bc = colors.whitesmoke
            else:  # Odd rows
                bc = colors.HexColor('#f8f9fa')
            
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, i), (-1, i), bc)
            ]))
        
        # Add table to story
        story.append(table)
        
        # Add page numbers and timestamp
        def add_page_number(canvas, doc):
            page_num = canvas.getPageNumber()
            text = f"Page {page_num}"
            canvas.saveState()
            canvas.setFont('Helvetica', 8)
            canvas.drawRightString(doc.width + doc.rightMargin - 20, 20, text)
            canvas.drawString(doc.leftMargin, 20, datetime.datetime.now().strftime("%Y-%m-%d %H:%M"))
            canvas.restoreState()
        
        # Build PDF with page numbers
        doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        import traceback
        st.error(f"Error generating ranking PDF: {str(e)}\n\n{traceback.format_exc()}")
        return None
    
def load_lottiefile(filepath: str):
    """ Loads a Lottie file (handles plain JSON and zipped .lottie files). """
    def read_json_from_stream(f_bytes_io, source_info=""):
        try: return json.loads(f_bytes_io.read().decode('utf-8'))
        except UnicodeDecodeError:
            try: f_bytes_io.seek(0); return json.loads(f_bytes_io.read().decode('cp1252'))
            except Exception as e_inner: st.error(f"Lottie Load Error: Failed decode ({source_info}): {e_inner}", icon="🚨"); return None
        except json.JSONDecodeError as e_json: st.error(f"Lottie Load Error: Invalid JSON ({source_info}): {e_json}", icon="🚨"); return None
        except Exception as e_read: st.error(f"Lottie Load Error: Read error ({source_info}): {e_read}", icon="🚨"); return None

    if not os.path.exists(filepath):
          return None
    try:
        with open(filepath, "rb") as f: file_bytes = f.read()
        if len(file_bytes) == 0:
            return None
        with io.BytesIO(file_bytes) as file_like_object:
            try: # Try ZIP
                with zipfile.ZipFile(file_like_object, "r") as z:
                    json_filename = None; namelist = z.namelist()
                    if "manifest.json" in namelist:
                        try:
                            with z.open("manifest.json") as mf: manifest = json.load(mf)
                            animations_list = manifest.get("animations")
                            if isinstance(animations_list, list) and animations_list:
                                first_animation = animations_list[0]
                                if isinstance(first_animation, dict): json_filename = first_animation.get("path")
                        except Exception as e_manifest: 
                            pass
                    if not json_filename: # Fallback search
                        known_paths = ["animations/data.json", "data.json", "animations/12345.json"]
                        for path in known_paths:
                            if path in namelist: json_filename = path; break
                        if not json_filename:
                             for name in namelist:
                                 if name.endswith('.json') and name != "manifest.json": json_filename = name; break
                    if json_filename:
                        try:
                            with z.open(json_filename) as jf: json_data = json.load(jf); return json_data
                        except Exception as e_open_json: 
                            return None
                    else: 
                        return None
            except zipfile.BadZipFile: # Try Plain JSON
                file_like_object.seek(0); json_data = read_json_from_stream(file_like_object, "plain JSON"); return json_data
            except Exception as e_zip: 
                return None
    except Exception as e_outer: 
        return None


def inject_particle_background():
    """ Injects a dynamic, interactive particle background using tsParticles. """
    st.components.v1.html(
        f"""
        <style>
        /* Make the streamlit app transparent so particle background is visible */
        .stApp {{
            background: transparent !important;
        }}
        /* Style for the particle container */
        #tsparticles {{
            position: fixed; top: 0; left: 0; width: 100%; height: 100%;
            z-index: -1; /* Send it to the back */
        }}
        </style>
        <div id="tsparticles"></div>
        <script src="https://cdn.jsdelivr.net/npm/tsparticles@3/tsparticles.bundle.min.js"></script>
        <script>
        (async () => {{{{
            await tsParticles.load({{{{
                id: "tsparticles",
                options: {{{{
                    "autoPlay": true,
                    "background": {{{{
                        "color": {{{{ "value": "#0d1117" }}}},
                        "image": "linear-gradient(-45deg, #0d1117, #161b22, #001f3f, #003366)",
                        "position": "50% 50%", "repeat": "no-repeat", "size": "cover", "opacity": 1
                    }}}},
                    "fullScreen": {{{{ "enable": false, "zIndex": -1 }}}},
                    "detectRetina": true, "fpsLimit": 60,
                    "interactivity": {{{{
                        "detectsOn": "window",
                        "events": {{{{
                            "onClick": {{{{ "enable": true, "mode": "push" }}}},
                            "onHover": {{{{
                                "enable": true, "mode": "bubble",
                                "parallax": {{{{ "enable": true, "force": 60, "smooth": 10 }}}}
                            }}}},
                            "resize": {{{{ "delay": 0.5, "enable": true }}}}
                        }}}},
                        "modes": {{{{
                            "bubble": {{{{
                                "distance": 200, "duration": 2, "opacity": 1, "size": 8
                            }}}},
                            "push": {{{{ "quantity": 4 }}}}
                        }}}}
                    }}}},
                    "motion": {{{{ "disable": false, "reduce": {{{{ "factor": 4, "value": true }}}} }}}},
                    "particles": {{{{
                        "color": {{{{ "value": "#ffffff" }}}},
                        "links": {{{{
                            "color": {{{{ "value": "#ffffff" }}}},
                            "distance": 150, "enable": true, "opacity": 0.4, "width": 1,
                            "shadow": {{{{
                                "enable": true,
                                "color": "#00c6ff",
                                "blur": 5
                            }}}}
                        }}}},
                        "move": {{{{
                            "direction": "none", "enable": true, "outModes": {{{{ "default": "out" }}}},
                            "random": false, "speed": 1.5, "straight": false, "vibrate": true
                        }}}},
                        "number": {{{{
                            "density": {{{{ "enable": true, "area": 800 }}}},
                            "value": 100
                        }}}},
                        "opacity": {{{{
                            "animation": {{{{ "enable": true, "speed": 1, "sync": false, "minimumValue": 0.1, "startValue": "random" }}}},
                            "value": 0.5
                        }}}},
                        "shape": {{{{ "type": "circle" }}}},
                        "size": {{{{
                            "animation": {{{{ "enable": true, "speed": 3, "sync": false, "minimumValue": 0.1, "startValue": "random" }}}},
                            "value": {{{{ "min": 1, "max": 3 }}}}
                        }}}}
                    }}}},
                    "pauseOnBlur": true, "pauseOnOutsideViewport": true
                }}}}
            }});
        }}}})();
        </script>
        """,
        height=0, width=0,
    )

def local_css(file_name):
    """ Loads a local CSS file and injects it into the app. """
    try:
        if not os.path.exists(file_name):
            return
        with open(file_name, encoding="utf-8") as f:
            css_content = f.read()
            if not css_content.strip():
                 pass
            st.markdown(f'<style>{css_content}</style>', unsafe_allow_html=True)
    except Exception as e_css:
         pass

# --- Prompts ---

cover_letter_prompt_template = """
Act as a professional career coach and expert resume writer.
Your task is to write a compelling, tailored cover letter for a job application.
The tone should be professional, confident, and enthusiastic.

CRITICAL TASK: Analyze the Job Description (JD) and the candidate's Resume Text. For the body paragraphs, you MUST identify 3-4 specific requirements in the JD and link them directly to concrete, quantifiable accomplishments or project descriptions (from the resume). Do not use general summaries.

- Use a standard 3-paragraph structure:
  1.  **Introduction:** State the role, where it was seen. Briefly express strong interest and state why they are a great fit (1-2 key skills).
  2.  **Body Paragraph(s):** Provide specific evidence. For each JD requirement, cite a specific project or experience bullet point and EXPLAIN how it meets the requirement.
  3.  **Conclusion:** Reiterate enthusiasm and include a call to action.
- **IMPORTANT:** Base all claims *only* on the provided resume text.
- Address "Dear Hiring Manager," and sign off "Sincerely,\n[Candidate Name]".
- Replace "[Candidate Name]" with the actual name from the resume data.
---
**Job Description:**
{jd}
---
**Candidate's Resume Text:**
{resume_text}
---
Now, please write the tailored cover letter.
"""

mock_interview_system_prompt = """
You are a strict but fair technical interviewer conducting a mock interview.
Based ONLY on the provided Job Description.
**Instructions:**
1.  **Start:** Introduce yourself briefly and immediately ask the first question.
2.  **Questions:** Ask a mix of technical and behavioral questions derived *directly* from the JD skills/requirements. Ask ONE question at a time and WAIT for the user's response.
3.  **Follow-up:** One brief follow-up if the answer is unclear, focused on the original question.
4.  **Feedback Trigger:** After exactly **5** user answers, STOP asking questions. Provide constructive feedback on their 5 answers (strengths/weaknesses vs. JD).
5.  **Stay in Character:** Professional and objective. No small talk.
---
**Job Description Context:**
{jd}
---
Okay, let's begin. Introduce yourself and ask the first question.
"""

skill_gap_learning_prompt_template = """
Act as an expert career coach and technical recruiter specializing in skill development.
Analyze the provided resume text against the target job description (JD).
**Your Task:**
Identify key skill gaps and suggest actionable learning steps. Format as markdown:
1.  **🎯 Top 3 Skill Gaps:**
    * List 2-3 *most significant* technical/tool skills from JD *missing* in resume.
    * Explain *why* each gap is significant for this role.
2.  **↔️ Potential Transferable Skills:**
    * Identify 1-2 skills *present* in resume that are transferable to JD needs.
    * Explain *how* (e.g., "Experience with [Resume Skill] applies to [JD Skill] because...").
3.  **📚 Actionable Learning Suggestions:**
    * For **each** 'Top 3 Skill Gap', provide 1-2 *specific* learning suggestions.
    * Examples: "Focus on [Topic]...", "Online courses on [Platform] covering '[Topic]'...", "Build a '[Project Type]' using [Skill]...", "Contribute to open-source...", "Practice on [Platform]..."
* Base analysis *only* on provided texts. Be realistic and encouraging.
---
**Job Description:**
{jd}
---
**Candidate's Resume Text:**
{resume_text}
---
Now, provide the skill gap analysis and learning suggestions in Markdown format.
"""

single_resume_prompt = """
Act as an elite ATS and expert tech recruiter. Analyze the resume against the JD.
**Resume Text:**
{text}
**Job Description (JD):**
{jd}
Provide output in a single, valid JSON object:
{{
  "JD_Match": "%",
  "MissingKeywords": ["...", "..."],
  "Matched_Keywords": ["...", "..."],
  "Profile_Summary": "Professional summary of candidate's fit.",
  "Analysis": {{
    "Quantitative_Check": [ {{ "Requirement": "...", "Found": true/false }} ],
    "Qualitative_Summary": "Short expert assessment of strengths/weaknesses."
  }},
  "Green_Flags": ["2-3 positive signals."],
  "Red_Flags": ["2-3 potential concerns."],
  "Interview_Questions": {{
    "Technical": ["..."],
    "Behavioral": ["..."]
  }},
  "Suggested_Resume_Summary": "A re-written, powerful summary aligned with JD.",
  "Skill_Map": {{
    "Programming_Languages": ["...", "..."],
    "Tools": ["...", "..."]
  }}
}}
"""

ranking_prompt_template = """
Act as a **Lead Algorithmic Recruiter** for a {company_name} hiring for a {job_title}.
Your analysis must be **dynamic**, meaning the scoring criteria MUST be derived exclusively from the provided Job Description (JD).

**CRITICAL DYNAMIC TASK:**
1.  **Define Criteria:** Analyze the JD and determine the specific technical and non-technical skills that correspond to the three score categories below.
2.  **Evaluate:** Score and rank all resumes based on this dynamically derived criteria.
3.  **Output:** Return a single, valid JSON list of objects, ranked from best to worst candidate.

**Job Description (JD):**
{jd}
---
**Here are the candidates:**
{candidate_texts}
---
**JSON Output Structure:**
[
  {{
    "Rank": 1,
    "Filename": "Candidate_A.pdf",
    "Overall_ATS_Score": "92%",
    "Confidence_Rating": "High - Excellent Technical Fit",
    "Scores_Breakdown": {{
      "Technical_Density": "X/5 (Measures mastery of core programming languages/frameworks required by the JD.)",
      "Experience_Alignment": "Y/5 (Measures relevance and tenure of past roles vs. the JD's seniority/domain.)",
      "Quantification_Metric": "Z/5 (Measures the use of measurable results and impact statements.)",
      "Experience_Benchmarking": "A/5 (Measures years of experience against the JD's stated/implied requirement.)",
      "Degree/Certification_Match": "B/5 (Measures relevance of education and required certifications.)"
    }},
    "Missing_Critical_Skills": ["Kubernetes (as required by JD)", "TypeScript (as required by JD)"],
    "Justification": "Brief summary of why this candidate is highly ranked, citing JD requirements."
  }}
]
"""

resume_quality_prompt_template = """
Act as a **Professional and Objective Career Strategist**. Your primary goal is to provide an **honest, evidence-based assessment** of the resume's readiness for a modern ATS and recruiter review. Frame all detailed critique constructively, focusing on **developmental steps** required to achieve an industry-standard score.

**CRITICAL INSTRUCTIONS:**
1.  **Date Verification:** Assume the current date is **{current_date}**. Any dates in the candidate's resume that are set after the current date MUST be flagged under the Red Flags section.
2.  **Focus of Critique:** Maintain the tone of an **Encouraging Mentor** ONLY when delivering the actionable suggestions, but remain the **Objective Expert** when assessing strengths, weaknesses, and assigning the score.
3.  **Enhanced Detail:** Provide a deep dive into Quantification and Strategic Integration (see JSON schema).

**Resume Text:**
{text}

Provide output in a single, valid JSON object. Ensure the arrays contain 2-4 detailed, specific, and crisp bullet points each for maximum value.

{{
  "Overall_Score": "xx/100",
  "ATS_Friendliness": {{
    "Score": "e.g., Excellent/Fair",
    "Explanation": "Detailed assessment of section flow, formatting, and keyword placement for ATS scanning. Be micro-detailed about layout risks (e.g., custom fonts, tables)."
  }},
  "Detailed_Improvement_Analysis": [
    {{ 
      "Area": "Section/Concept to improve (e.g., Professional Summary, Experience Bullet Points, Skills Formatting)", 
      "Rating": "e.g., Strong, Needs Polish, Needs Development",
      "Micro_Problems": [
        "**Specific, low-level problem 1.** (e.g., 30% of bullets start with 'Responsible for', diluting impact).", 
        "**Specific, low-level problem 2.** (e.g., Dates for 'Company X' are misaligned and may confuse ATS timeline parsing).",
        "**Specific, low-level problem 3.** (e.g., The Skills section is a block of text, not a clear list, making skill extraction difficult).",
        "**Specific, low-level problem 4.** (e.g., Achievements lack quantifiable impact: the 'Increased X' is missing the 'by Y%' metric)."
      ],
      "Executive_Suggestions": [
        "**Clear, crisp action 1.** (e.g., Rewrite summary to start with: 'Senior [Role] with 8+ years...').", 
        "**Clear, crisp action 2.** (e.g., Apply the X-Y-Z formula to all experience points: 'Accomplished [X] as measured by [Y] by doing [Z]').",
        "**Clear, crisp action 3.** (e.g., Group skills into categories (e.g., Languages, Frameworks, Tools) to improve readability and ATS parsing).",
        "**Clear, crisp action 4.** (e.g., Ensure all dates are in consistent format (e.g., MM/YYYY - MM/YYYY) for clean chronological ordering)."
      ]
    }} 
  ],
  "Strategic_Review": {{
    "Quantification_Score": "X/5 (Assessment of metric usage and impact.)",
    "Impact_Verb_Ratio": "Y% (e.g., 40%. Measures the ratio of active/quantifiable verbs used in experience.)",
    "Salary_Estimate_Bracket": "Estimate the salary bracket (e.g., $60K-$80K USD or Mid-Level Range) based ONLY on the skills and experience provided in the resume.",
    "Integration_Notes": "Assessment of how effectively skills are woven into experience narratives, not just listed. Suggest 2-3 advanced integrations."
  }},
  "Red_Flags": ["2-3 potential concerns, including any dates after {current_date}."],
  "Improved_Resume_Text": "Full rewritten resume text as a single, cleaned string, incorporating all suggested changes into the structure for maximum impact."
}}
"""

predictive_qa_prompt_template = """
Act as an experienced senior hiring manager preparing to interview a candidate.
Given: Candidate's Resume, Job Description (JD).
Task: Generate 5-7 highly specific, personalized interview questions based on *connections* and *gaps* between these two. Do NOT ask generic questions.
Focus on 3 categories:
1.  **Resume-to-JD Probes (2-3 questions):** Find item on **resume**. Find related requirement in **JD**. Ask question *connecting* them. *Example:* "I see 'Project X' using Python on your resume. The JD emphasizes data analysis. Walk me through a data cleaning challenge in that project."
2.  **Behavioral/Situational Questions (2-3 questions):** Find soft skill/responsibility in **JD** (e.g., "teamwork"). Find context on **resume** (e.g., "Team Lead at Club Y"). Ask "Tell me about a time..." question using resume context. *Example:* "The JD mentions collaboration. I noticed you were 'Lead' for 'Project Y'. Tell me about a time you resolved a conflict within that team."
3.  **Potential Gap/Weakness Probes (1-2 questions):** Find key requirement in **JD** (e.g., "AWS") **NOT** on **resume**. Ask direct, fair question. *Example:* "The role requires 'AWS'. Your resume shows strong development, but I didn't see 'AWS' experience. Tell me about any exposure you've had to cloud platforms."
---
**Job Description:**
{jd}
---
**Candidate's Resume Text:**
{resume_text}
---
Now, generate the 5-7 personalized questions in clear markdown, grouped by category.
"""

book_recommendation_prompt_template = """
Act as an expert book recommendation engine. You have received a list of skills the user needs to learn.
**Your Task:** For each skill deficiency, generate the most relevant, highly-reviewed book or course title and a brief 2-sentence justification for its selection, simulating a ranking and clustering process.

**Skills Deficiencies:**
{missing_skills_list}

Provide output in a clear Markdown list format:
1.  **[Skill Name]:** [Book Title] by [Author]. *Justification: [Briefly explain the book's relevance to the skill gap and its 'high review score'].*
2.  
"""

portfolio_generator_prompt = """
Act as a **Staff Engineer and Lead UX Strategist**. Your task is to generate a **Cinematic, High-Performance, and Recruiter-Ready** portfolio website. The aesthetic must be **Cyber-Minimalist**: highly functional, dark-themed, and exceptionally clear.

Generate a **complete, single-file, responsive** portfolio website.
- CSS MUST be in a <style> tag.
- All JavaScript MUST be in a <script> tag.
- ALL external links MUST open in a new tab (`target="_blank"`).
- Use the provided colors: **Primary Accent: Electric Blue (#00BFFF)**, **Secondary Accent: Cyan (#00E0FF)**, **Deep Background: Dark Charcoal (#121212)**.

**RECRUITER-FOCUSED UI/UX REQUIREMENTS:**

1.  **Aesthetic Foundation (Minimalist Dark):** Use the deep background color. Containers must use subtle **flat drop shadows and border gradients**—**NO NEUMORPHISM**. The focus is on clean lines and high contrast.
    * **Chromatic Aberration:** Apply a subtle `filter: drop-shadow()` effect using the cyan and Electric Blue tones to headings on hover.

2.  **Hero Section & Geometry (CRITICAL CONSOLIDATION):**
    * **Structure:** The main hero section (`<section id="hero">`) MUST contain the **Typewriter Name/Title**. The **Geometry Animation (`tsparticles`)** must be integrated as the hero section's background, appearing in one unified area.
    * **Hero Text (Name):** Text color MUST be **Pure White (`#FFFFFF`)**. Shadow effect must be a **Single, Subtle Electric Blue Glow (#00BFFF)**. NO Glitch effect.
    * **Functional Movement:** The accompanying titles MUST use the **JS Typewriter Effect**.
    * **Navigation Structure (ELIMINATED):** **STRICTLY DO NOT RENDER ANY VISIBLE NAVBAR, HEADER LINKS, OR TOP TITLES.**
    * **Scroll Indicator:** Header MUST include a subtle **Electric Blue Scroll-Progress Indicator Bar** at the top.

3.  **Skills Section (VISUAL EFFECT & LABEL FIX):**
    * **Label Fix:** The Skills section heading MUST be labeled **'Skills'** (using `<section id="skills">`).
    * **Visual Effect:** Individual skill tags/chips MUST have a **soft Electric Blue Glow-on-Hover** effect (`box-shadow` or similar).

4.  **Project Cards (CRITICAL CONDITIONAL REDIRECTION):**
    * **Link Guarantee:** The external project link button/anchor MUST be present and clickable (opening in a **new tab**) **ONLY IF** the `link` field in the JSON data is present and valid. If the link is missing or empty, the button must be omitted or visually disabled (e.g., "Link N/A").
    * **Interaction:** On hover, the project card MUST use **Slight upward translation (`transform: translateY(-8px)`) and a visible border highlight using the Electric Blue Accent**.

5.  **Motion (Clean Animation - Guaranteed):** Sections must use the **AOS/IntersectionObserver** pattern for entrance animations, demanding a **clean, quick fade-in/slide-up**.

6.  **Particles/Geometry (CRITICAL ANIMATION - GUARANTEED):** Integrate `tsparticles` with low count (e.g., 30-50), slow speed, and **mandatory visible Electric Blue connection lines between particles**.

7.  **Contact Section (SIMPLIFIED FOOTER LOGIC):**
    * The contact footer (`<section id="contact">`) MUST clearly display the user's **Email** and **LinkedIn** link.
    * **Conditional Resume Download:** A prominent "Download Resume" button MUST be rendered **ONLY IF** the `resume_data_uri` is present (not empty).

8.  **GLOBAL ALIGNMENT & SPACING (CRITICAL FIX):** All main text paragraphs (especially in the "About" section and under project descriptions) MUST have:
    * **Tighter Line Height:** A reduced line height (e.g., `line-height: 1.5;`) to close up vertical gaps.
    * **Tighter Vertical Margins:** Reduced top/bottom margins on paragraph and section elements (e.g., `margin-bottom: 0.8em;`).
    * **Proper Alignment:** All primary text, including the "About" section summary, MUST be **left-aligned** (`text-align: left;`) for optimal readability.

---
**DATA (as a JSON object):**
{data_json}
**RESUME (as a Base64 Data URI):**
{resume_data_uri}
---
Generate the complete, single-file HTML. Start with `<!DOCTYPE html>` and end with `</html>`. **The output must be pure, working HTML/CSS/JS.** Do NOT wrap code in markdown backticks.
"""

resume_parser_for_portfolio_prompt = """
Act as an expert data extraction bot.
Analyze resume text, extract info in structured JSON.
Output MUST be a single, valid JSON object.
**Resume Text:**
{text}
**JSON Output Structure:**
{{
  "name": "Full Name",
  "email": "email@example.com",
  "linkedin": "linkedin.com/in/...",
  "github": "github.com/...",
  "summary": "The professional summary...",
  "skills": "Comma, separated, list, of, skills",
  "projects": [
    {{ "name": "Project 1", "description": "1-2 lines.", "link": "https://github..." }}
  ]
}}
**Instructions:**
1. Extract accurately. If not found, return "" (strings) or [] (lists).
2. For `skills`, combine into single comma-separated string.
3. For `projects`, find name, description, and URL. If no link, "link": "".
"""

linkedin_optimization_prompt = """
Act as an Elite LinkedIn Optimization Engine. You will analyze a user's LinkedIn profile PDF text.

Your goals:
1. Rewrite their headline for maximum recruiter visibility.
2. Rewrite their About section for clarity and impact.
3. Rewrite weak Experience bullets into strong measurable achievements.
4. Identify missing high-value keywords for their industry.
5. Compute a LinkedIn Profile Completeness Score (0–100).
6. Compute a Search Ranking Score (0–100) based on keyword match & density.
7. Provide an actionable improvement plan.

INPUT (LinkedIn Profile Extracted Text):
{resume_text}

OUTPUT (STRICT VALID JSON):
{
  "Optimal_Headline": "string",
  "Optimized_Summary": "string",
  "Experience_Rewrite": [
    {
      "Original": "string",
      "Improved": "string"
    }
  ],
  "Missing_Profile_Keywords": ["keyword1", "keyword2"],
  "Keyword_Density": {
    "Strong": ["word1","word2"],
    "Weak": ["word3","word4"],
    "Missing": ["word5","word6"]
  },
  "Completeness_Score": "85/100",
  "Search_Ranking_Score": "78/100",
  "Behavioral_Strategy": {
    "Visual_Showcase_Suggestions": ["tip1", "tip2"],
    "Call_to_Action_Strategy": "string"
  }
}
"""


# --- Streamlit App UI ---
st.set_page_config(layout="wide", page_title=" hiREsume 📄🎯")

# --- Inject Particle Background ---
inject_particle_background()

# --- Load CSS ---
local_css("style.css") 

# --- Initialize Session State ---
if 'experience' not in st.session_state: st.session_state.experience = []
if 'education' not in st.session_state: st.session_state.education = []
if 'projects' not in st.session_state: st.session_state.projects = []
if 'achievements' not in st.session_state: st.session_state.achievements = []
if 'certificates' not in st.session_state: st.session_state.certificates = []
if 'custom_links' not in st.session_state: st.session_state.custom_links = []
if 'docx_buffer' not in st.session_state: st.session_state.docx_buffer = None
if 'pdf_buffer' not in st.session_state: st.session_state.pdf_buffer = None
if 'resume_filename' not in st.session_state: st.session_state.resume_filename = "resume"
if 'quality_response' not in st.session_state: st.session_state.quality_response = ""
if 'improved_resume_text' not in st.session_state: st.session_state.improved_resume_text = ""
if 'quality_filename' not in st.session_state: st.session_state.quality_filename = ""
if 'generated_cover_letter' not in st.session_state: st.session_state.generated_cover_letter = ""
if 'mock_interview_messages' not in st.session_state: st.session_state.mock_interview_messages = []
if 'mock_interview_jd' not in st.session_state: st.session_state.mock_interview_jd = ""
if 'mock_interview_error' not in st.session_state: st.session_state.mock_interview_error = None
if 'skill_gap_analysis_output' not in st.session_state: st.session_state.skill_gap_analysis_output = ""
if 'jd_analyzer_output' not in st.session_state: st.session_state.jd_analyzer_output = ""
if 'trend_jds' not in st.session_state: st.session_state.trend_jds = []
if 'trend_analysis_output' not in st.session_state: st.session_state.trend_analysis_output = ""
if 'predictive_qa_output' not in st.session_state: st.session_state.predictive_qa_output = ""
if 'portfolio_projects' not in st.session_state: st.session_state.portfolio_projects = []
if 'portfolio_generated_code' not in st.session_state: st.session_state.portfolio_generated_code = ""
if 'pf_name' not in st.session_state: st.session_state.pf_name = st.session_state.get('resume_name', '')
if 'pf_email' not in st.session_state: st.session_state.pf_email = st.session_state.get('resume_email', '')
if 'pf_linkedin' not in st.session_state: st.session_state.pf_linkedin = st.session_state.get('resume_linkedin', '')
if 'pf_github' not in st.session_state: st.session_state.pf_github = st.session_state.get('resume_github', '')
if 'pf_summary' not in st.session_state: st.session_state.pf_summary = st.session_state.get('resume_summary', '')
if 'pf_skills' not in st.session_state: st.session_state.pf_skills = st.session_state.get('resume_skills', '')
if 'pf_hero_text' not in st.session_state: st.session_state.pf_hero_text = ""
if 'active_view' not in st.session_state: st.session_state.active_view = "✍️ Resume Maker"
if 'resume_source_cl' not in st.session_state: st.session_state.resume_source_cl = "Upload a New Resume"
if 'resume_source_qa' not in st.session_state: st.session_state.resume_source_qa = "Upload a New Resume"
if 'api_response_tab1' not in st.session_state: st.session_state.api_response_tab1 = ""
if 'pdf_display_tab1' not in st.session_state: st.session_state.pdf_display_tab1 = ""
if 'resume_text_tab1' not in st.session_state: st.session_state.resume_text_tab1 = ""
if 'jd_text_tab1' not in st.session_state: st.session_state.jd_text_tab1 = ""
if 'uploaded_filename_tab1' not in st.session_state: st.session_state.uploaded_filename_tab1 = ""
if 'jd_multi' not in st.session_state: st.session_state.jd_multi = ""
if "chat_messages" not in st.session_state: st.session_state.chat_messages = []


# --- Sidebar Navigation ---
with st.sidebar:
    lottie_sidebar = None; lottie_load_error = None
    try:
        lottie_sidebar = load_lottiefile("STUDENT.lottie")
        if lottie_sidebar is None and os.path.exists("Loading Files.lottie"): lottie_load_error = "Animation load error."
        elif not os.path.exists("Loading Files.lottie"): lottie_load_error = "'Loading Files.lottie' not found."
    except Exception as e: lottie_load_error = f"Exception: {e}"
    if lottie_sidebar: st_lottie(lottie_sidebar, height=150, key="sidebar_lottie")
    elif lottie_load_error: st.warning(f"⚠️ Sidebar Anim Load Failed: {lottie_load_error}", icon="🎨")
    else: st.warning("⚠️ Sidebar animation load failed.", icon="🎨")
    st.markdown("<h1 style='text-align: center; font-size: 28px; color: #FFFFFF !important; text-shadow: 0 0 10px #00BFFF, 0 0 20px #00E0FF; font-weight: 700; '>hiREsume 📄</h1>", unsafe_allow_html=True)
    st.caption("<p style='text-align: center; font-size: 14px; font-style: italic;'>The Algorithm to Bypass Every ATS</p>", unsafe_allow_html=True)
    st.divider()
    st.markdown("""
<h3 style='
    font-size: 32px; 
    font-weight: 700;
    text-align: center;
    padding-top: 15px; /* Adds space above the header */
    padding-bottom: 8px; /* Separates it cleanly from the divider */
    color: #FFFFFF;
'>
    🧭 NAVIGATOR
</h3>
""", unsafe_allow_html=True)
    def nav_button(label):
        is_active = (st.session_state.active_view == label)
        if st.button(label, use_container_width=True, type="primary" if is_active else "secondary"):
            st.session_state.active_view = label; st.rerun()
    with st.expander("📝 **CREATE DOCUMENTS**", expanded=True):
        nav_button("✍️ Resume Maker"); nav_button("✉️ AI Cover Letter Generator"); nav_button("🌐 AI Portfolio Generator")
    with st.expander("🔍 **ANALYZE & RANK**", expanded=True):
        nav_button("📄 Resume Quality Checker"); nav_button("🔗 LinkedIn Optimizer"); nav_button("👤 Single Resume Analysis"); nav_button("📊 Multi-Candidate Ranking")
        
    with st.expander("🔍 **SEARCH**", expanded=True):
        nav_button("🔍 Job Search")
    with st.expander("🚀 **INTERVIEW PREP**", expanded=True):
        nav_button("🤖 Mock Interviewer"); nav_button("🎯 Skill Gap Analyzer"); nav_button("🔮 Predictive Q&A")
    with st.expander("💬 **UTILITIES**", expanded=True): nav_button("💬 HR Chatbot")
    st.divider(); st.caption("Built with Streamlit & Google Gemini"); st.caption("Pro UI/UX Design by Gemini ✨")


# --- Main App Area ---
st.title(" hiREsume 📄")
active_view = st.session_state.active_view

# --- Resume Maker View ---
if active_view == "✍️ Resume Maker":
    st.header("✍️ ATS-Friendly Resume Maker")
    st.info("Fill details below for an ATS-friendly resume (.docx or .pdf).")
    st.subheader("1. PERSONAL INFORMATION")
    col1, col2 = st.columns(2)
    with col1:
        st.text_input("Full Name", placeholder="Krsna", key="resume_name")
    with col2:
        st.text_input("Location", placeholder="Bengaluru, India", key="resume_location")
    col3, col4 = st.columns(2)
    with col3:
        st.text_input("Email", placeholder="example@gmail.com", key="resume_email")
    with col4:
        st.text_input("Phone", placeholder="+91 1234567890", key="resume_phone")
    st.text_input("LinkedIn URL", placeholder="linkedin.com/in/your-profile", key="resume_linkedin")
    st.text_input("GitHub URL (Optional)", placeholder="github.com/your-username", key="resume_github")
    # Section 2: Custom Links
    st.subheader("2. CUSTOM LINKS")
    with st.form("links_form", clear_on_submit=True):
        link_label = st.text_input("Link Label", placeholder="e.g., Portfolio")
        link_url = st.text_input("Link URL or Text", placeholder="e.g., my-portfolio.com")
        submitted_link = st.form_submit_button("➕ Add Link")
        if submitted_link and link_label and link_url:
            st.session_state.custom_links.append({"label": link_label, "url": link_url}); st.success("Link added!")
        elif submitted_link: st.warning("Please fill in both Label and URL.")
    if st.session_state.custom_links:
        st.write("---"); st.write("**Current Custom Links:**")
        indices_to_remove_links = []
        for i, link in enumerate(st.session_state.custom_links):
            with st.container(border=True):
                st.markdown(f"**{link['label']}**: {link['url']}")
                if st.button("Remove 🗑️", key=f"rem_link_{i}", type="secondary"): indices_to_remove_links.append(i)
        if indices_to_remove_links:
             for index in sorted(indices_to_remove_links, reverse=True): del st.session_state.custom_links[index]
             st.rerun()
    # Section 3: Summary and Skills
    st.subheader("3. SUMMARY AND SKILLS")
    
    # --- Consolidated Form for Summary and Skills (Fixes 'Enter' issue) ---
    with st.form("summary_skills_form", clear_on_submit=False):
        # We use unique keys for the input widgets here
        st.text_area("Professional Summary", placeholder="A brief summary...", height=150, key="resume_summary_input", 
                     value=st.session_state.get('resume_summary', ''))
        st.text_area("Skills (Comma-separated)", placeholder="Python, Java, SQL...", key="resume_skills_input", 
                     value=st.session_state.get('resume_skills', ''))
        
        submitted_ss = st.form_submit_button("➕ Save Summary & Skills")
        
        # When the dedicated button is clicked, update the main session state variables
        if submitted_ss:
            st.session_state.resume_summary = st.session_state.resume_summary_input
            st.session_state.resume_skills = st.session_state.resume_skills_input
            st.success("Summary and Skills saved!", icon="✅")

    # Section 4: Education
    st.subheader("4. EDUCATION")
    with st.form("education_form", clear_on_submit=True):
        edu_degree = st.text_input("Degree"); edu_institution = st.text_input("Institution")
        edu_dates = st.text_input("Dates"); edu_details = st.text_area("Details (Optional)", placeholder="e.g., CGPA")
        submitted_edu = st.form_submit_button("➕ Add Education")
        if submitted_edu and all([edu_degree, edu_institution, edu_dates]):
            st.session_state.education.append({"degree": edu_degree, "institution": edu_institution, "dates": edu_dates, "details": edu_details}); st.success("Education added!")
        elif submitted_edu: st.warning("Please fill Degree, Institution, and Dates.")
    if st.session_state.education:
        st.write("---"); st.write("**Current Education Entries:**")
        indices_to_remove_edu = []
        for i, edu in enumerate(st.session_state.education):
            with st.container(border=True):
                st.markdown(f"**{edu['degree']}** from {edu['institution']} ({edu['dates']})")
                if st.button("Remove 🗑️", key=f"rem_edu_{i}", type="secondary"): indices_to_remove_edu.append(i)
        if indices_to_remove_edu:
             for index in sorted(indices_to_remove_edu, reverse=True): del st.session_state.education[index]
             st.rerun()
    # Section 5: Work Experience
    st.subheader("5. WORK EXPERIENCE")
    with st.form("experience_form", clear_on_submit=True):
        exp_title = st.text_input("Job Title"); exp_company = st.text_input("Company")
        exp_dates = st.text_input("Dates (e.g., Jan 2024 - Present)")
        exp_desc = st.text_area("Description / Accomplishments", height=150, help="One accomplishment per line. Use action verbs and quantify results (e.g., 'Optimized... increasing speed by 15%').")
        submitted_exp = st.form_submit_button("➕ Add Experience")
        if submitted_exp and all([exp_title, exp_company, exp_dates, exp_desc]):
            st.session_state.experience.append({"title": exp_title, "company": exp_company, "dates": exp_dates, "description": exp_desc}); st.success("Experience added!")
        elif submitted_exp: st.warning("Please fill all fields.")
    if st.session_state.experience:
        st.write("---"); st.write("**Current Experience Entries:**")
        indices_to_remove_exp = []
        for i, exp in enumerate(st.session_state.experience):
            with st.container(border=True):
                st.markdown(f"**{exp['title']}** at {exp['company']} ({exp['dates']})")
                if st.button("Remove 🗑️", key=f"rem_exp_{i}", type="secondary"): indices_to_remove_exp.append(i)
        if indices_to_remove_exp:
             for index in sorted(indices_to_remove_exp, reverse=True): del st.session_state.experience[index]
             st.rerun()
    # Section 6: Projects
    st.subheader("6. PROJECTS")
    with st.form("project_form", clear_on_submit=True):
        proj_name = st.text_input("Project Name"); proj_desc = st.text_area("Description", height=100, help="One point per line. Include link/stack.")
        submitted_proj = st.form_submit_button("➕ Add Project")
        if submitted_proj and all([proj_name, proj_desc]):
            st.session_state.projects.append({"name": proj_name, "description": proj_desc}); st.success("Project added!")
        elif submitted_proj: st.warning("Please fill all fields.")
    if st.session_state.projects:
        st.write("---"); st.write("**Current Project Entries:**")
        indices_to_remove_proj = []
        for i, proj in enumerate(st.session_state.projects):
            with st.container(border=True):
                st.markdown(f"**{proj['name']}**")
                if st.button("Remove 🗑️", key=f"rem_proj_{i}", type="secondary"): indices_to_remove_proj.append(i)
        if indices_to_remove_proj:
             for index in sorted(indices_to_remove_proj, reverse=True): del st.session_state.projects[index]
             st.rerun()
    # Section 7: Achievements
    st.subheader("7. ACHIEVEMENTS")
    with st.form("achievements_form", clear_on_submit=True):
        ach_name = st.text_input("Achievement"); submitted_ach = st.form_submit_button("➕ Add Achievement")
        if submitted_ach and ach_name:
            st.session_state.achievements.append({"name": ach_name}); st.success("Achievement added!")
        elif submitted_ach: st.warning("Please fill in the name.")
    if st.session_state.achievements:
        st.write("---"); st.write("**Current Achievements:**")
        indices_to_remove_ach = []
        for i, ach in enumerate(st.session_state.achievements):
            with st.container(border=True):
                st.markdown(f"• {ach['name']}")
                if st.button("Remove 🗑️", key=f"rem_ach_{i}", type="secondary"): indices_to_remove_ach.append(i)
        if indices_to_remove_ach:
             for index in sorted(indices_to_remove_ach, reverse=True): del st.session_state.achievements[index]
             st.rerun()
    # Section 8: Certificates
    st.subheader("8. CERTIFICATES")
    with st.form("certificates_form", clear_on_submit=True):
        cert_name = st.text_input("Certificate Name")
        cert_issuer = st.text_input("Issuer (Optional)")
        cert_date = st.text_input("Date (e.g., 2023)", placeholder="YYYY")
        submitted_cert = st.form_submit_button("➕ Add Certificate")
        
        if submitted_cert and cert_name:
            cert_data = {"name": cert_name}
            if cert_issuer:
                cert_data["issuer"] = cert_issuer
            if cert_date:
                cert_data["date"] = cert_date
            st.session_state.certificates.append(cert_data)
            st.rerun()
        elif submitted_cert:
            st.warning("Please enter at least a certificate name")

    # Display current certificates
    if st.session_state.certificates:
        st.write("---")
        st.write("**Current Certificates:**")
        indices_to_remove_cert = []
        
        for i, cert in enumerate(st.session_state.certificates):
            with st.container(border=True):
                col1, col2 = st.columns([5, 1])
                with col1:
                    st.markdown(f"**{cert.get('name', 'N/A')}**")
                    details = []
                    if cert.get('issuer'):
                        details.append(f"Issuer: {cert['issuer']}")
                    if cert.get('date'):
                        details.append(f"Date: {cert['date']}")
                    if details:
                        st.caption(" | ".join(details))
                with col2:
                    if st.button("🗑️", key=f"rem_cert_{i}"):
                        indices_to_remove_cert.append(i)
                        st.rerun()
        
        if indices_to_remove_cert:
            for index in sorted(indices_to_remove_cert, reverse=True):
                del st.session_state.certificates[index]
            st.rerun()
    
    st.divider()
    st.subheader("9. GENERATE YOUR RESUME")
    validation_messages = []; is_valid = True; current_resume_name = st.session_state.get('resume_name', '')
    required_fields = ['resume_name', 'resume_email', 'resume_phone', 'resume_location', 'resume_summary', 'resume_skills']
    required_lists = {'experience': 'Work Experience', 'education': 'Education'}
    for field in required_fields:
        if not st.session_state.get(field): validation_messages.append(f"🚨 Please fill in '{field.replace('resume_', '').replace('_', ' ').title()}'."); is_valid = False
    for list_key, list_name in required_lists.items():
          if not st.session_state.get(list_key): validation_messages.append(f"🚨 Please add at least one '{list_name}' entry."); is_valid = False
    resume_data = { "name": current_resume_name, "email": st.session_state.get('resume_email', ''), "phone": st.session_state.get('resume_phone', ''),
                     "linkedin": st.session_state.get('resume_linkedin', ''), "github": st.session_state.get('resume_github', ''),
                     "custom_links": st.session_state.custom_links, "location": st.session_state.get('resume_location', ''),
                     "summary": st.session_state.get('resume_summary', ''), "skills": st.session_state.get('resume_skills', ''),
                     "experience": st.session_state.experience, "education": st.session_state.education,
                     "projects": st.session_state.projects, "achievements": st.session_state.achievements,
                     "certificates": st.session_state.certificates }
    
    gen_col1, gen_col2 = st.columns(2)
    with gen_col1:
        if st.button("Generate .docx 📄", use_container_width=True, type="primary", key="gen_docx"):
            st.session_state.pdf_buffer = None; st.session_state.docx_buffer = None
            if is_valid:
                placeholder_docx = st.empty(); lottie_data = None
                try: lottie_data = load_lottiefile("Loading Files.lottie")
                except Exception as e: st.warning(f"Anim load failed: {e}")
                if lottie_data:
                     with placeholder_docx: st_lottie(lottie_data, height=200, key="lottie_docx_gen")
                else:
                     with placeholder_docx: st.info("📄 Generating DOCX...")
                try: buffer = create_resume_docx(resume_data)
                finally: placeholder_docx.empty()
                if buffer: st.session_state.docx_buffer = buffer; st.session_state.resume_filename = f"{current_resume_name.replace(' ', '_')}_Resume"; st.rerun()
                else: st.error("Failed to generate .docx document.")
            else:
                for msg in validation_messages: st.error(msg)
    with gen_col2:
        if st.button("Generate .pdf 📑", use_container_width=True, type="primary", key="gen_pdf"):
            st.session_state.pdf_buffer = None; st.session_state.docx_buffer = None
            if is_valid:
                placeholder_pdf = st.empty(); lottie_data = None
                try: lottie_data = load_lottiefile("Loading Files.lottie")
                except Exception as e: st.warning(f"Anim load failed: {e}")
                if lottie_data:
                     with placeholder_pdf: st_lottie(lottie_data, height=200, key="lottie_pdf_gen")
                else:
                     with placeholder_pdf: st.info("📑 Generating PDF...")
                try: buffer = create_resume_pdf(resume_data)
                finally: placeholder_pdf.empty()
                if buffer: st.session_state.pdf_buffer = buffer; st.session_state.resume_filename = f"{current_resume_name.replace(' ', '_')}_Resume"; st.rerun()
                else: st.error("Failed to generate .pdf document.")
            else:
                for msg in validation_messages: st.error(msg)
    
    st.divider()
    if st.session_state.docx_buffer or st.session_state.pdf_buffer:
        st.subheader("🎉 Success! What's next?")
        col1, col2, col3 = st.columns([1, 1, 1.2])
        with col1:
            if st.session_state.docx_buffer:
                st.download_button(label="⬇️ Download .docx", data=st.session_state.docx_buffer, file_name=f"{st.session_state.resume_filename}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            else: st.button("⬇️ Download .docx", use_container_width=True, disabled=True, help="Generate .docx first.")
        with col2:
            if st.session_state.pdf_buffer:
                st.download_button(label="⬇️ Download .pdf", data=st.session_state.pdf_buffer, file_name=f"{st.session_state.resume_filename}.pdf", mime="application/pdf", use_container_width=True)
            else: st.button("⬇️ Download .pdf", use_container_width=True, disabled=True, help="Generate .pdf first.")
        with col3:
            st.write("**Smart Actions:**")
            col3a, col3b = st.columns(2)
            with col3a:
                if st.button("✉️ Gen Cover Letter", use_container_width=True, type="secondary", help="Use this resume data."):
                    st.session_state.active_view = "✉️ AI Cover Letter Generator"; st.session_state.resume_source_cl = "Use Data from Resume Maker (Tab 1)"; st.rerun()
            with col3b:
                if st.button("🔮 Prep Interview", use_container_width=True, type="secondary", help="Use this resume data."):
                    st.session_state.active_view = "🔮 Predictive Q&A"; st.session_state.resume_source_qa = "Use Data from Resume Maker (Tab 1)"; st.rerun()

# --- Resume Quality Checker View ---
elif active_view == "📄 Resume Quality Checker":
    st.header("📄 Resume Quality Checker & Improver")
    st.info("Upload resume for score, **micro-detailed** feedback, and rewritten version.")
    uploaded_file_quality = st.file_uploader("Upload Your Resume", type=["pdf", "docx"], key="uploader_quality")
    submit_quality = st.button("🔬 Analyze & Improve Resume", key="submit_quality", type="primary")
    if submit_quality:
        st.session_state.quality_response = ""
        st.session_state.improved_resume_text = ""
        st.session_state.quality_filename = ""
        process_error = None
        
        if uploaded_file_quality is None:
            st.error("🚨 Please upload a resume file.")
        else:
            loading_placeholder_quality = st.empty()
            lottie_data = None
            try:
                lottie_data = load_lottiefile("Scanning Document.lottie")
            except Exception as e:
                st.warning(f"Could not load animation: {e}", icon="⚠️")
            
            if lottie_data:
                with loading_placeholder_quality:
                    st_lottie(lottie_data, height=300, key="lottie_quality_scan")
            else:
                with loading_placeholder_quality:
                    st.info("🤖 Scanning resume...")
                    
            try:
                original_filename = uploaded_file_quality.name
                base_filename = ".".join(original_filename.split('.')[:-1])
                st.session_state.quality_filename = f"{base_filename}_Improved_Resume.txt"
                text = extract_text_from_file(uploaded_file_quality)
                
                if text:
                    # --- DYNAMIC DATE CALCULATION AND INJECTION ---
                    current_date_str = datetime.datetime.now().strftime("%B %Y")
                    
                    # Use .format() to inject both the resume text and the dynamic date string
                    formatted_prompt = resume_quality_prompt_template.format(
                        text=text,
                        current_date=current_date_str
                    )
                    
                    # Call API and save response
                    response_text = get_gemini_response(formatted_prompt, temperature=0.0)
                    st.session_state.quality_response = response_text
                    
                    if response_text.startswith("Error"):
                        process_error = response_text
                else:
                    process_error = "❌ Failed to extract text."
                    st.session_state.quality_response = ""
                    
            except Exception as e_process:
                process_error = f"Error: {e_process}"
                st.session_state.quality_response = ""
            
            finally:
                loading_placeholder_quality.empty()

        # ... (The remaining UI display logic follows this block)
            
        if process_error and (not st.session_state.quality_response or not st.session_state.quality_response.startswith("Error")): 
            st.error(process_error, icon="🚨")
        elif not process_error and st.session_state.quality_response and not st.session_state.quality_response.startswith("Error"): 
            st.success("Analysis complete!", icon="✅")

    if st.session_state.quality_response:
        # ... (The remaining UI display logic continues here)
        response_text = st.session_state.quality_response
        if response_text.startswith("Error"): st.error(f"Analysis Error: {response_text}", icon="🚨")
        else:
             try:
                 clean_json_text = clean_json_response(response_text)
                 response_json = json.loads(clean_json_text)
                 score_str = response_json.get('Overall_Score', '0/100')
                 try: score_value = int(score_str.split('/')[0])
                 except Exception: score_value = 0
                 ats_friendliness = response_json.get('ATS_Friendliness', {})
                 improvements = response_json.get('Detailed_Improvement_Analysis', [])
                 st.session_state.improved_resume_text = response_json.get('Improved_Resume_Text', 'Error: Could not rewrite resume text.')
                 
                 # Define base_filename here to resolve scope error
                 base_filename = uploaded_file_quality.name.split('.')[:-1][0] if uploaded_file_quality else "Improved_Resume" 

                 # Pull strategic review data safely
                 strategic_review = response_json.get('Strategic_Review', {})
                 salary_estimate = strategic_review.get('Salary_Estimate_Bracket', 'N/A')

                 st.subheader("📊 Quality Analysis")
                 col1, col2, col3 = st.columns(3)
                 col1.metric("Overall Score", score_str, delta=f"{score_value-75}" if score_value > 0 else None, help="AI's estimate of 'job-readiness'.")
                 col2.metric("ATS Friendliness", ats_friendliness.get('Score', 'N/A'), help=ats_friendliness.get('Explanation', 'No explanation.'))
                 col3.metric("Estimated Salary Range", salary_estimate, help="AI-modeled income based on skills and experience.")
                 st.info(f"**ATS Breakdown:** {ats_friendliness.get('Explanation', 'N/A')}")
                 st.divider()

                 # Strategic / Impact metrics (if available)
                 st.markdown("##### 🎯 Impact & Integration Metrics")
                 try:
                     st.metric("Quantification Score", strategic_review.get('Quantification_Score', 'N/A'))
                 except Exception:
                     # Fallback in case strategic_review is not a mapping
                     st.metric("Quantification Score", 'N/A')
                 st.caption(f"**Impact Verb Ratio:** {strategic_review.get('Impact_Verb_Ratio', 'N/A')}")
                 st.markdown(f"**Integration Notes:** {strategic_review.get('Integration_Notes', 'N/A')}")

                 st.subheader("✨ Your Improved Resume Text")
                 st.info("Copy text into Word/Google Doc, fix formatting, then save as PDF/DOCX.")
                 
                 # --- UPDATED: Resume Quality Download Buttons (PARALLEL FIX) ---
                 col_dl_1, col_dl_2 = st.columns(2)
                 
                 docx_buffer = create_markdown_docx(st.session_state.improved_resume_text, title="Improved Resume Text")
                 with col_dl_1:
                     if docx_buffer: st.download_button(label="⬇️ Download Improved (.docx)", data=docx_buffer, file_name=f"{base_filename}_Improved.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, key="dl_improved_docx")
                     else: st.button("Error DOCX", disabled=True, use_container_width=True, key="error_improved_docx")
                 
                 # PDF for Improved Resume (using the markdown converter)
                 pdf_buffer_improved = create_markdown_pdf(st.session_state.improved_resume_text, title="Improved Resume Text")
                 with col_dl_2:
                     if pdf_buffer_improved: st.download_button(label="⬇️ Download Improved (.pdf)", data=pdf_buffer_improved, file_name=f"{base_filename}_Improved.pdf", mime="application/pdf", use_container_width=True, key="dl_improved_pdf")
                     else: st.button("Error PDF", disabled=True, use_container_width=True, key="error_improved_pdf")
                 # --- END UPDATED DOWNLOADS ---

                 with st.expander("View new text"): st.text(st.session_state.improved_resume_text)
                 st.divider()
                 
                 st.subheader("💡 Micro-Detailed Improvements")
                 if improvements:
                      for item in improvements:
                          with st.expander(f"**{item.get('Area', 'General Area')}** - *Rating: {item.get('Rating', 'N/A')}*"):
                              st.markdown("##### ❌ Micro-Problems:")
                              for prob in item.get('Micro_Problems', []): st.write(f"- {prob}")
                              st.markdown("##### 🚀 Executive Suggestions:")
                              for sugg in item.get('Executive_Suggestions', []): st.write(f"- {sugg}")
                 else: st.success("No major micro-level improvements suggested. This is an excellent resume!", icon="✅")

             except json.JSONDecodeError as e: st.error(f"Error parsing AI response: {e}", icon="🚨"); st.text("Raw Response:"); st.code(response_text, language=None)
             except Exception as e_display: st.error(f"Error displaying results: {e_display}", icon="🚨"); st.text("Raw Response:"); st.code(response_text, language=None)

# --- AI Cover Letter Generator View ---
elif active_view == "✉️ AI Cover Letter Generator":
    st.header("✉️ AI Cover Letter Generator")
    st.info("Writes a cover letter tailored to a job, using your resume data.")
    jd_text_cl = st.text_area("Paste the Job Description here", height=250, key="jd_cl")
    st.subheader("Select Your Resume Source")
    resume_options = ["Use Data from Resume Maker (Tab 1)", "Upload a New Resume"]
    try: current_index_cl = resume_options.index(st.session_state.resume_source_cl)
    except ValueError: current_index_cl = 1
    resume_source = st.radio("Resume data source:", resume_options, key="resume_source_cl_radio",
                             index=current_index_cl, horizontal=True,
                             on_change=lambda: st.session_state.update(resume_source_cl=st.session_state.resume_source_cl_radio))
    uploaded_file_cl = None
    if resume_source == "Upload a New Resume":
        uploaded_file_cl = st.file_uploader("Upload Your Resume", type=["pdf", "docx"], key="uploader_cl")
    if st.button("🚀 Generate Cover Letter", type="primary", use_container_width=True, key="submit_cl"):
        st.session_state.generated_cover_letter = ""; resume_text = None; api_error = None; resume_load_error = None
        if not jd_text_cl: st.error("🚨 Please paste the Job Description first.")
        elif resume_source == "Upload a New Resume" and uploaded_file_cl is None: st.error("🚨 Please upload a resume file.")
        elif resume_source == "Use Data from Resume Maker (Tab 1)" and not st.session_state.get('experience'): st.error("🚨 'Resume Maker' is empty. Fill it out or upload a resume.")
        else:
            loading_placeholder_cl = st.empty()
            lottie_data = None; 
            try: lottie_data = load_lottiefile("Loading Files.lottie") # Default
            except Exception as e: st.warning(f"Could not load animation: {e}", icon="⚠️")
            if lottie_data:
                 with loading_placeholder_cl: st_lottie(lottie_data, height=300, key="lottie_cl_write")
            else:
                 with loading_placeholder_cl: st.info("✍️ Generating cover letter...")
            try:
                if resume_source == "Use Data from Resume Maker (Tab 1)":
                    resume_data_cl = { "name": st.session_state.get('resume_name', ''), "email": st.session_state.get('resume_email', ''),
                                       "phone": st.session_state.get('resume_phone', ''), "linkedin": st.session_state.get('resume_linkedin', ''),
                                       "github": st.session_state.get('resume_github', ''), "custom_links": st.session_state.custom_links,
                                       "location": st.session_state.get('resume_location', ''), "summary": st.session_state.get('resume_summary', ''),
                                       "skills": st.session_state.get('resume_skills', ''), "experience": st.session_state.experience,
                                       "education": st.session_state.education, "projects": st.session_state.projects,
                                       "achievements": st.session_state.achievements }
                    resume_text = format_resume_data_for_prompt(resume_data_cl)
                    if not resume_text: resume_load_error = "❌ Error formatting data from Resume Maker."
                else: # Upload a New Resume
                    resume_text = extract_text_from_file(uploaded_file_cl)
                    if not resume_text: resume_load_error = "❌ Failed to extract text from the uploaded resume."
                if resume_text and not resume_load_error: # Only call API if resume text is valid
                    formatted_prompt = cover_letter_prompt_template.format(jd=jd_text_cl, resume_text=resume_text)
                    response_text = get_gemini_response(formatted_prompt, temperature=0.5)
                    if response_text.startswith("Error"): api_error = response_text
                    else: st.session_state.generated_cover_letter = response_text
            except Exception as e_process: api_error = f"An unexpected error occurred: {e_process}"; st.session_state.generated_cover_letter = ""
            finally: loading_placeholder_cl.empty()
            if resume_load_error: st.error(resume_load_error, icon="🚨")
            elif api_error: st.error(api_error, icon="🚨")
            elif st.session_state.generated_cover_letter: st.success("Cover letter generated!", icon="✅")
            
    # Display results
    if st.session_state.generated_cover_letter:
        st.divider(); st.subheader("Your Generated Cover Letter")
        st.text_area("Copy your letter from here:", value=st.session_state.generated_cover_letter, height=400, key="cl_output_area")
        st.write("**Download your letter:**"); col1, col2 = st.columns(2)
        with col1:
            docx_buffer = create_cover_letter_docx(st.session_state.generated_cover_letter)
            if docx_buffer: st.download_button(label="⬇️ Download .docx", data=docx_buffer, file_name="Cover_Letter.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, key="download_docx_cl")
            else: st.button("Error DOCX", disabled=True, use_container_width=True, key="error_docx_cl")
        with col2:
            pdf_buffer = create_cover_letter_pdf(st.session_state.generated_cover_letter)
            if pdf_buffer: st.download_button(label="⬇️ Download .pdf", data=pdf_buffer, file_name="Cover_Letter.pdf", mime="application/pdf", use_container_width=True, key="download_pdf_cl")
            else: st.button("Error PDF", disabled=True, use_container_width=True, key="error_pdf_cl")

# --- Mock Interviewer View ---
elif active_view == "🤖 Mock Interviewer":
    st.header("🤖 Mock Interview Simulator")
    st.info("Paste a Job Description below and click 'Start Interview' to practice.")
    jd_text_mock = st.text_area("Paste the Job Description here", height=250, key="jd_mock")
    if st.button("Start / Restart Interview", type="primary", key="start_mock"):
        st.session_state.mock_interview_messages = []; st.session_state.mock_interview_error = None
        start_error = None
        if not jd_text_mock: st.warning("Please paste the Job Description first.")
        else:
            st.session_state.mock_interview_jd = jd_text_mock
            initial_prompt_text = mock_interview_system_prompt.format(jd=jd_text_mock)
            initial_context = [{'role': 'user', 'parts': [{'text': initial_prompt_text}]}]
            loading_placeholder_mock = st.empty()
            lottie_data = None
            try: lottie_data = load_lottiefile("Loading Files.lottie") # Default, or use 'interview.lottie'
            except Exception as e: st.warning(f"Could not load animation: {e}", icon="⚠️")
            if lottie_data:
                 with loading_placeholder_mock: st_lottie(lottie_data, height=300, key="lottie_mock_start")
            else:
                 with loading_placeholder_mock: st.info("🎙️ Preparing interview...")
            try:
                stream = get_gemini_response_chat(initial_context, temperature=0.7)
                full_response = "";
                for chunk in stream:
                    if hasattr(chunk, 'text') and chunk.text: full_response += chunk.text
                if "Error" in full_response or "blocked" in full_response.lower() or not full_response.strip():
                    start_error = "Failed to start interview. API Error: " + full_response
                    st.session_state.mock_interview_jd = ""
                else:
                    st.session_state.mock_interview_messages.append({"role": "assistant", "content": full_response})
            except Exception as e:
                start_error = f"An error occurred while starting the interview: {e}"
                st.session_state.mock_interview_jd = ""
            finally: loading_placeholder_mock.empty()
            if start_error: st.error(start_error, icon="🚨")
            elif st.session_state.mock_interview_messages: st.rerun()
    st.divider()
    # --- Display Chat History ---
    for message in st.session_state.mock_interview_messages:
        with st.chat_message(message["role"]): st.markdown(message["content"])
    # --- User Input ---
    if prompt := st.chat_input("Your answer...", key="mock_input"):
        if not st.session_state.mock_interview_jd: st.warning("Please start an interview first.", icon="⚠️")
        else:
            st.session_state.mock_interview_messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"): st.markdown(prompt)
            chat_context = [{'role': 'user', 'parts': [{'text': mock_interview_system_prompt.format(jd=st.session_state.mock_interview_jd)}]}]
            for msg in st.session_state.mock_interview_messages:
                role = 'user' if msg['role'] == 'user' else 'model'; chat_context.append({'role': role, 'parts': [{'text': msg['content']}]})
            user_message_count = sum(1 for msg in st.session_state.mock_interview_messages if msg["role"] == "user")
            if user_message_count == 5:
                chat_context.append({'role': 'user', 'parts': [{'text': "Provide feedback now."}]}); st.info("🏁 Generating feedback...")
            with st.chat_message("assistant"):
                message_placeholder = st.empty(); full_response = ""
                try:
                    stream = get_gemini_response_chat(chat_context, temperature=0.7)
                    for chunk in stream:
                        if hasattr(chunk, 'text') and chunk.text:
                            full_response += chunk.text; message_placeholder.markdown(full_response + "▌")
                    message_placeholder.markdown(full_response)
                    if "Error" in full_response or "blocked" in full_response.lower() or not full_response.strip():
                        st.error("Error receiving response.", icon="🚨")
                    else: st.session_state.mock_interview_messages.append({"role": "assistant", "content": full_response})
                except Exception as e:
                    full_response = f"An unexpected error occurred: {e}"; message_placeholder.error(full_response, icon="🚨")
                    st.session_state.mock_interview_messages.append({"role": "assistant", "content": full_response})

# --- Skill Gap Analyzer View ---
elif active_view == "🎯 Skill Gap Analyzer":
    st.header("🎯 Skill Gap Analyzer & Learning Suggester")
    st.info("Upload resume and paste JD to identify skill gaps.")
    uploaded_file_sg = st.file_uploader("Upload Your Resume", type=["pdf", "docx"], key="uploader_sg")
    jd_text_sg = st.text_area("Paste the Target Job Description here", height=250, key="jd_sg")
    if st.button("🔬 Analyze Skill Gap & Get Suggestions", type="primary", use_container_width=True, key="submit_sg"):
        st.session_state.skill_gap_analysis_output = ""; analysis_error = None
        if uploaded_file_sg is None: st.error("🚨 Please upload your resume file.")
        elif not jd_text_sg: st.error("🚨 Please paste the Job Description.")
        else:
            loading_placeholder_sg = st.empty()
            lottie_data = None
            try: lottie_data = load_lottiefile("ai ai.lottie") # Use specific Lottie
            except Exception as e: st.warning(f"Could not load animation: {e}", icon="⚠️")
            if lottie_data:
                 with loading_placeholder_sg: st_lottie(lottie_data, height=200, key="lottie_sg_load")
            else:
                 with loading_placeholder_sg: st.info("🧠 Analyzing skill gaps...")
            try:
                resume_text_sg = extract_text_from_file(uploaded_file_sg)
                if not resume_text_sg: analysis_error = "❌ Failed to extract text from your resume."
                else:
                    formatted_prompt = skill_gap_learning_prompt_template.format(jd=jd_text_sg, resume_text=resume_text_sg)
                    response_text = get_gemini_response(formatted_prompt, temperature=0.0)
                    if response_text.startswith("Error"): analysis_error = response_text
                    else: st.session_state.skill_gap_analysis_output = response_text
            except Exception as e_process: analysis_error = f"An error occurred: {e_process}"
            finally: loading_placeholder_sg.empty()
            if analysis_error: st.session_state.skill_gap_analysis_output = ""; st.error(analysis_error, icon="🚨")
            elif st.session_state.skill_gap_analysis_output: st.success("Analysis complete!", icon="✅")
            
    if st.session_state.skill_gap_analysis_output:
        st.divider(); st.subheader("💡 Skill Gap Analysis & Learning Suggestions")
        st.markdown(st.session_state.skill_gap_analysis_output)

        # --- NEW: Book Recommendation Module ---
        st.divider()
        st.subheader("📚 Strategic Learning Path")

        # Extract the missing skills from the analysis text
        # (Simulating extraction if the LLM output is not JSON for this report)
        missing_skills_text = re.search(r'🎯 Top 3 Skill Gaps:\s*\n[\s\S]*?(?=2\. \u2194)', st.session_state.skill_gap_analysis_output)
        missing_skills_list = missing_skills_text.group(0).split('\n')[1:] if missing_skills_text else ["Critical skills missing"]

        if missing_skills_list and missing_skills_list[0].strip() not in ["", "Critical skills missing"]:
            with st.spinner("📚 Simulating learning path recommendation..."):
                book_prompt = book_recommendation_prompt_template.format(missing_skills_list="\n".join(missing_skills_list))
                book_response = get_gemini_response(book_prompt, temperature=0.3)
            
                if book_response.startswith("Error"):
                    st.warning("Could not generate book recommendations due to API error.", icon="⚠️")
                else:
                    st.markdown(book_response)
        else:
            st.info("No critical skill gaps were identified, so no learning path is needed!")

        st.divider()
        # --- UPDATED: Skill Gap Download Buttons (DOCX + PDF - PARALLEL FIX) ---
        col1, col2 = st.columns(2)

        docx_buffer = create_markdown_docx(st.session_state.skill_gap_analysis_output, title="Skill Gap Analysis Report")
        with col1:
            if docx_buffer: st.download_button(label="⬇️ Download Report (.docx)", data=docx_buffer, file_name="Skill_Gap_Analysis.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, key="download_sg_docx")
            else: st.button("Error DOCX", disabled=True, use_container_width=True, key="error_sg_docx")
        
        pdf_buffer = create_markdown_pdf(st.session_state.skill_gap_analysis_output, title="Skill Gap Analysis Report")
        with col2:
            if pdf_buffer: st.download_button(label="⬇️ Download Report (.pdf)", data=pdf_buffer, file_name="Skill_Gap_Analysis.pdf", mime="application/pdf", use_container_width=True, key="download_sg_pdf")
            else: st.button("Error PDF", disabled=True, use_container_width=True, key="error_sg_pdf")


# --- Single Resume Analysis View ---
elif active_view == "👤 Single Resume Analysis":
    st.header("Analyze a Single Candidate (vs. Job Description)")
    jd_tab1_input = st.text_area("Paste the Job Description", height=200, key="jd_input_tab1")
    uploaded_file_tab1 = st.file_uploader("Upload Resume", type=["pdf", "docx"], help="Supports text, scanned PDF (OCR), and DOCX", key="uploader_tab1")
    submit_tab1 = st.button("🔍 Analyze Resume", key="submit_tab1", type="primary")
    if submit_tab1:
        st.session_state.api_response_tab1 = ""; st.session_state.pdf_display_tab1 = ""
        st.session_state.resume_text_tab1 = ""; st.session_state.jd_text_tab1 = ""; file_bytes = None
        st.session_state.uploaded_filename_tab1 = ""; analysis_error = None
        if uploaded_file_tab1 is None: st.error("🚨 Please upload a resume file.")
        elif not jd_tab1_input: st.error("📋 Please paste the job description.")
        else:
            loading_placeholder_single = st.empty()
            lottie_data = None
            try: lottie_data = load_lottiefile("Data Scanning.lottie") # Use specific Lottie
            except Exception as e: st.warning(f"Could not load animation: {e}", icon="⚠️")
            if lottie_data:
                 with loading_placeholder_single: st_lottie(lottie_data, height=300, key="lottie_single_scan")
            else:
                 with loading_placeholder_single: st.info("📊 Analyzing resume against JD...")
            try:
                st.session_state.jd_text_tab1 = jd_tab1_input
                st.session_state.uploaded_filename_tab1 = uploaded_file_tab1.name
                file_bytes = uploaded_file_tab1.getvalue()
                text = extract_text_from_file(uploaded_file_tab1)
                if text:
                    st.session_state.resume_text_tab1 = text
                    if uploaded_file_tab1.name.endswith('.pdf'):
                        try:
                            base64_pdf = base64.b64encode(file_bytes).decode('utf-8')
                            st.session_state.pdf_display_tab1 = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="700" style="border: none;"></iframe>'
                        except Exception as e_pdf:
                            st.warning(f"PDF preview error: {e_pdf}", icon="⚠️")
                            st.session_state.pdf_display_tab1 = "PDF preview failed."
                    else: st.session_state.pdf_display_tab1 = "📄 Preview not available for .docx."
                    formatted_prompt = single_resume_prompt.format(text=text, jd=st.session_state.jd_text_tab1)
                    response_text = get_gemini_response(formatted_prompt, temperature=0.0)
                    st.session_state.api_response_tab1 = response_text
                    if response_text.startswith("Error"): analysis_error = response_text
                else: analysis_error = "❌ Failed to extract text."
            except Exception as e_process:
                 analysis_error = f"An unexpected error occurred: {e_process}"; st.session_state.api_response_tab1 = ""
            finally: loading_placeholder_single.empty()
            if analysis_error and (not st.session_state.api_response_tab1 or not st.session_state.api_response_tab1.startswith("Error")):
                 st.error(analysis_error, icon="🚨")
            elif not analysis_error and st.session_state.api_response_tab1: st.success("Analysis complete!", icon="✅")
            
    # Display results
    if st.session_state.api_response_tab1:
        col1, col2 = st.columns([2, 3])
        with col1:
            st.subheader("📄 Resume Viewer")
            if st.session_state.pdf_display_tab1.startswith('<iframe'): st.markdown(st.session_state.pdf_display_tab1, unsafe_allow_html=True)
            elif st.session_state.resume_text_tab1 :
                 with st.expander("Show Extracted Text", expanded=False): st.text(st.session_state.resume_text_tab1)
            else: st.info(st.session_state.pdf_display_tab1 if st.session_state.pdf_display_tab1 else "No preview available.")
        
        with col2:
            st.subheader("🤖 AI Analysis Dashboard")
            response_text = st.session_state.api_response_tab1
            if response_text.startswith("Error"): st.error(f"Analysis Error: {response_text}", icon="🚨")
            else:
                 try:
                      clean_json_text = clean_json_response(response_text)
                      response_json = json.loads(clean_json_text)
                      match_percent_str = response_json.get('JD_Match', '0%')
                      match_value = int(re.sub(r'[^0-9]', '', match_percent_str) or 0)
                      missing_keywords = response_json.get('MissingKeywords', [])
                      matched_keywords = response_json.get('Matched_Keywords', [])
                      
                      st.markdown("##### Core Metrics")
                      metric_cols = st.columns(3)
                      metric_cols[0].metric(label="🎯 ATS Score", value=f"{match_value}%", delta=f"{match_value-50}%" if match_value else None, help="AI's estimate of match.")
                      metric_cols[1].metric(label="✅ Keywords Found", value=len(matched_keywords))
                      metric_cols[2].metric(label="❌ Keywords Missing", value=len(missing_keywords))
                      st.divider()
                      
                      st.markdown("**📝 Profile Summary:**"); st.write(response_json.get('Profile_Summary', 'N/A'))
                      
                      with st.expander("👁️ Show Keyword Match Highlighter"):
                           hl_cols = st.columns(2)
                           with hl_cols[0]: st.markdown("**Job Description**"); annotated_text(build_annotated_text(st.session_state.jd_text_tab1, matched_keywords))
                           with hl_cols[1]: st.markdown("**Resume Text**"); annotated_text(build_annotated_text(st.session_state.resume_text_tab1, matched_keywords))
                      
                      with st.expander("🚦 Show Recruiter Flags (Red/Green)"):
                           green_flags = response_json.get('Green_Flags', []); red_flags = response_json.get('Red_Flags', [])
                           st.markdown("**✅ Green Flags (Positives):**");
                           if green_flags:
                               for flag in green_flags: st.write(f"- {flag}")
                           else: st.caption("None identified.")
                           st.markdown("**🚩 Red Flags (Concerns):**");
                           if red_flags:
                               for flag in red_flags: st.write(f"- {flag}")
                           else: st.caption("None identified.")
                      
                      with st.expander("❓ Show Suggested Interview Questions"):
                           questions = response_json.get('Interview_Questions', {}); tech_q = questions.get('Technical', []); behav_q = questions.get('Behavioral', [])
                           st.markdown("**💻 Technical Questions:**")
                           if tech_q:
                               for i, q in enumerate(tech_q, 1): st.write(f"{i}. {q}")
                           else: st.caption("None suggested.")
                           st.markdown("**🗣️ Behavioral Questions:**")
                           if behav_q:
                               for i, q in enumerate(behav_q, 1): st.write(f"{i}. {q}")
                           else: st.caption("None suggested.")
                      
                      # --- UPDATED: Single Resume Analysis Download Buttons (DOCX + PDF - PARALLEL FIX) ---
                      st.divider()
                      st.markdown("##### Download Full Report")
                      
                      report_docx_data = create_analysis_report_docx(response_json, st.session_state.uploaded_filename_tab1)
                      dl_filename = re.sub(r'[^a-zA-Z0-9_.-]', '_', st.session_state.uploaded_filename_tab1)
                      
                      col_dl_1, col_dl_2 = st.columns(2)

                      with col_dl_1:
                          if report_docx_data:
                              st.download_button(label="⬇️ Download Report (.docx)", data=report_docx_data, file_name=f"{dl_filename}_ATS_Analysis.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, key="download_analysis_docx")
                          else:
                              st.button("Error DOCX", disabled=True, use_container_width=True, key="error_analysis_docx")

                      # PDF generation for this specific JSON-based report
                      report_markdown_text = json.dumps(response_json, indent=2) # Use the raw JSON as markdown text for conversion
                      report_pdf_data = create_markdown_pdf(report_markdown_text, title=f"ATS Analysis for {st.session_state.uploaded_filename_tab1}")

                      with col_dl_2:
                          if report_pdf_data:
                              st.download_button(label="⬇️ Download Report (.pdf)", data=report_pdf_data, file_name=f"{dl_filename}_ATS_Analysis.pdf", mime="application/pdf", use_container_width=True, key="download_analysis_pdf")
                          else:
                              st.button("Error PDF", disabled=True, use_container_width=True, key="error_analysis_pdf")
                      # --- END UPDATED DOWNLOADS ---

                 except json.JSONDecodeError as e: st.error(f"Error parsing AI response: {e}", icon="🚨"); st.text("Raw Response:"); st.code(response_text, language=None)
                 except Exception as e_display: st.error(f"Error displaying results: {e_display}", icon="🚨"); st.text("Raw Response:"); st.code(response_text, language=None)

# --- Multi-Candidate Ranking View ---
elif active_view == "📊 Multi-Candidate Ranking":
    st.header("📊 Rank Multiple Candidates (vs. Job Description)")
    jd_multi_input = st.text_area("Paste the Job Description", height=200, key="jd_input_multi")
    uploaded_files_multi = st.file_uploader("Upload Resumes (Max 10 per batch)", type=["pdf", "docx"], accept_multiple_files=True, key="uploader_multi")
    rank_button = st.button("📊 Rank Candidates", key="submit_multi", type="primary")

    if rank_button:
        # Initialize variables
        st.session_state.jd_multi = jd_multi_input
        ranking_df = None  # Crucial: Initialize outside the try block
        ranking_error = None
        files_skipped = 0

        if not uploaded_files_multi:
            st.error("🚨 Please upload at least one resume.")
        elif not st.session_state.jd_multi:
            st.error("📋 Please paste the job description.")
        else:
            loading_placeholder_multi = st.empty()
            lottie_data = None
            
            try:
                lottie_data = load_lottiefile("Ai-powered marketing tools abstract.lottie")
            except Exception as e:
                st.warning(f"Could not load animation: {e}", icon="⚠️")
                
            if lottie_data:
                # NOTE: Using unique key is important here
                loading_lottie_key = f"lottie_multi_rank_load_{datetime.datetime.now().timestamp()}"
                with loading_placeholder_multi: st_lottie(lottie_data, height=200, key=loading_lottie_key)
            else:
                with loading_placeholder_multi: st.info("📈 Ranking candidates...")
            
            try:
                candidate_data = []
                for file in uploaded_files_multi:
                    text = extract_text_from_file(file)
                    if text: candidate_data.append((file.name, text))
                    else: files_skipped += 1
                    
                if candidate_data:
                    candidate_texts_for_prompt = "\n\n".join([f"---\nCandidate Filename: {fname}\nResume Text:\n{ftext}\n---" for fname, ftext in candidate_data])
                    jd_lines = st.session_state.jd_multi.split('\n')
                    job_title = jd_lines[0].strip() if jd_lines else "role"
                    company_name = next((l.split(":",1)[1].strip() for l in jd_lines[:5] if ":" in l and ("company" in l.lower() or "organization" in l.lower())), "your company")
                    
                    ranking_prompt = ranking_prompt_template.format(company_name=company_name, job_title=job_title, jd=st.session_state.jd_multi, candidate_texts=candidate_texts_for_prompt)
                    response_text = get_gemini_response(ranking_prompt, temperature=0.0)
                    
                    if response_text.startswith("Error"): 
                        ranking_error = response_text
                    else:
                        try:
                            clean_json_text = clean_json_response(response_text)
                            ranking_json = json.loads(clean_json_text)
                            df = pd.DataFrame(ranking_json)
                            # Handle column renaming for compatibility
                            if "Overall_ATS_Score" in df.columns: df = df.rename(columns={"Overall_ATS_Score": "Score"})
                            if "Confidence_Rating" in df.columns: df = df.rename(columns={"Confidence_Rating": "Confidence"})
                            
                            ranking_df = df
                        except Exception as parse_e: 
                            ranking_error = f"Error parsing ranking: {parse_e}\nRaw:{response_text}"
                else: 
                    ranking_error = "Could not read text from any uploaded files."
            
            except Exception as e_process: 
                ranking_error = f"An unexpected error occurred: {e_process}"
            
            finally: 
                loading_placeholder_multi.empty()
                
            # --- DISPLAY LOGIC (This block now executes outside the inner try/except) ---
            if ranking_error: 
                st.error(ranking_error, icon="🚨")
            
            # --- Success Display ---
            elif ranking_df is not None:
                st.subheader("🏆 Candidate Ranking & Comparative Analytics")
                
                # Convert DataFrame for main display
                st.dataframe(ranking_df[['Rank', 'Filename', 'Score', 'Confidence', 'Justification']], 
                             use_container_width=True, hide_index=True)
                
                st.divider()
                st.subheader("📊 Detailed Comparative Scorecards")
                
                # Iterate through the DataFrame rows for expanded view
                for index, row in ranking_df.iterrows():
                    rank = row['Rank']
                    filename = row['Filename']
                    overall_score = row['Score']
                    confidence = row['Confidence']
                    justification = row['Justification']
                    
                    # Access the nested JSON fields (Safely checking for existence)
                    scores = row.get('Scores_Breakdown', {})
                    missing_skills = row.get('Missing_Critical_Skills', [])
                    
                    # Display each candidate's detailed scorecard in an expander
                    with st.expander(f"Rank {rank}: {filename} ({overall_score}) - Confidence: {confidence}", expanded=(rank==1)):
                        
                        # --- Score Breakdown ---
                        st.markdown("##### Performance Breakdown")
                        col_metrics = st.columns(3)
                        
                        # Extract the data and remove the parentheses text for clean display
                        tech_score = scores.get("Technical_Density", "N/A").split(' ')[0]
                        align_score = scores.get("Experience_Alignment", "N/A").split(' ')[0]
                        quant_score = scores.get("Quantification_Metric", "N/A").split(' ')[0]

                        # Display Scores
                        col_metrics[0].metric("Technical Density", tech_score)
                        col_metrics[1].metric("Experience Alignment", align_score)
                        col_metrics[2].metric("Quantification Metric", quant_score)
                        
                        # Display Explanations below for clarity (New Section)
                        st.markdown("---") 
                        st.markdown("**Criteria Defined:**")
                        st.caption(f"**Technical Density:** {scores.get('Technical_Density', 'N/A').split('(', 1)[-1].strip(')')}")
                        st.caption(f"**Experience Alignment:** {scores.get('Experience_Alignment', 'N/A').split('(', 1)[-1].strip(')')}")
                        st.caption(f"**Quantification Metric:** {scores.get('Quantification_Metric', 'N/A').split('(', 1)[-1].strip(')')}")
                        st.markdown("---")
                        
                        # --- Missing Skills ---
                        st.markdown("##### ❌ Critical Skill Gaps")
                        if missing_skills:
                            # Ensure output is clean, remove template guidance if present
                            clean_skills = [s.split('(')[0].strip() for s in missing_skills]
                            st.warning(f"**Missing:** {', '.join(clean_skills)}", icon="⚠️")
                        else:
                            st.success("No critical gaps identified.")
                        
                        # --- Recruiter Justification ---
                        st.markdown("##### Recruiter Justification")
                        st.write(justification)


                st.success("Ranking and analysis complete!", icon="✅")
                
                st.markdown("##### Download Reports")
                col_dl_1, col_dl_2 = st.columns(2)

                # 1. CSV Download (Retained)
                csv = ranking_df.to_csv(index=False).encode('utf-8')
                with col_dl_1:
                    st.download_button("⬇️ Download Ranking as CSV", data=csv, file_name="candidate_ranking_analysis.csv", mime="text/csv", use_container_width=True, key="dl_ranking_csv")

                # 2. PDF Download (New Feature)
                pdf_buffer = create_ranking_pdf(ranking_df, job_title=job_title)
                with col_dl_2:
                    if pdf_buffer:
                        st.download_button("⬇️ Download Ranking as PDF", data=pdf_buffer, file_name="candidate_ranking_analysis.pdf", mime="application/pdf", use_container_width=True, key="dl_ranking_pdf")
                    else:
                        st.button("Error PDF", disabled=True, use_container_width=True, key="error_ranking_pdf")
            
            # Display file skip warning if necessary
            if files_skipped > 0:
                st.warning(f"Note: {files_skipped} file(s) skipped due to extraction errors.", icon="⚠️")

# --- End of Multi-Candidate Ranking View ---

# --- End of Multi-Candidate Ranking View ---

# --- HR Chatbot View ---
elif active_view == "💬 HR Chatbot":
    st.header("💬 HR Assistant Chatbot")
    st.info("Ask HR questions, get interview tips, or ask about the loaded JD.")
    for message in st.session_state.chat_messages:
        with st.chat_message(message["role"]): st.markdown(message["content"])
    active_jd = st.session_state.get('jd_text_tab1', "") or st.session_state.get('jd_multi', "")
    if active_jd: st.caption(f"ℹ️ Using JD context: `{active_jd[:100]}...`")
    if prompt := st.chat_input("Ask the HR Assistant..."):
        st.session_state.chat_messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"): st.markdown(prompt)
        with st.chat_message("assistant"):
            message_placeholder = st.empty(); full_response = ""
            chat_context = [{'role': 'user', 'parts': [{'text': "You are a helpful HR Assistant."}]}, {'role': 'model', 'parts': [{'text': "Okay, I'm ready."}]}]
            if active_jd:
                chat_context.append({'role': 'user', 'parts': [{'text': f"Use this JD context if relevant:\n{active_jd}"}]})
                chat_context.append({'role': 'model', 'parts': [{'text': "Okay, I have the JD context."}]})
            history_start_index = max(0, len(st.session_state.chat_messages) - 11)
            for msg in st.session_state.chat_messages[history_start_index:]:
                role = 'user' if msg['role'] == 'user' else 'model'; chat_context.append({'role': role, 'parts': [{'text': msg['content']}]})
            try:
                stream = get_gemini_response_chat(chat_context, temperature=0.7)
                for chunk in stream:
                    if hasattr(chunk, 'text') and chunk.text:
                        full_response += chunk.text; message_placeholder.markdown(full_response + "▌")
                message_placeholder.markdown(full_response)
                if "Error" in full_response or "blocked" in full_response.lower() or not full_response.strip():
                    st.error("Error receiving response.", icon="🚨")
                else: st.session_state.chat_messages.append({"role": "assistant", "content": full_response})
            except Exception as e:
                full_response = f"An unexpected error occurred: {e}"; message_placeholder.error(full_response, icon="🚨")
                st.session_state.chat_messages.append({"role": "assistant", "content": full_response})


# --- Predictive Q&A View ---
elif active_view == "🔮 Predictive Q&A":
    st.header("🔮 Predictive Interview Q&A Generator")
    st.info("Get personalized questions based on resume and target JD.")
    jd_text_qa = st.text_area("Paste the Job Description here", height=250, key="jd_qa")
    st.subheader("Select Your Resume Source")
    resume_options_qa = ["Use Data from Resume Maker (Tab 1)", "Upload a New Resume"]
    try: current_index_qa = resume_options_qa.index(st.session_state.resume_source_qa)
    except ValueError: current_index_qa = 1
    resume_source_qa = st.radio("Resume data source:", resume_options_qa, key="resume_source_qa_radio",
                                 index=current_index_qa, horizontal=True,
                                 on_change=lambda: st.session_state.update(resume_source_qa=st.session_state.resume_source_qa_radio))
    uploaded_file_qa = None
    if resume_source_qa == "Upload a New Resume":
        uploaded_file_qa = st.file_uploader("Upload Your Resume", type=["pdf", "docx"], key="uploader_qa")
    if st.button("🚀 Generate Predicted Questions", type="primary", use_container_width=True, key="submit_qa"):
        st.session_state.predictive_qa_output = ""; resume_text = None; analysis_error = None
        if not jd_text_qa: st.error("🚨 Please paste the Job Description first.")
        elif resume_source_qa == "Upload a New Resume" and uploaded_file_qa is None: st.error("🚨 Please upload a resume file.")
        elif resume_source_qa == "Use Data from Resume Maker (Tab 1)" and not st.session_state.get('experience'): st.error("🚨 'Resume Maker' is empty. Fill it out or upload a resume.")
        else:
            loading_placeholder_qa = st.empty()
            lottie_data = None
            try: lottie_data = load_lottiefile("Loading Files.lottie") # Default
            except Exception as e: st.warning(f"Could not load animation: {e}", icon="⚠️")
            if lottie_data:
                 with loading_placeholder_qa: st_lottie(lottie_data, height=300, key="lottie_qa_predict")
            else:
                 with loading_placeholder_qa: st.info("🔮 Generating questions...")
            try:
                if resume_source_qa == "Use Data from Resume Maker (Tab 1)":
                    resume_data_qa = { "name": st.session_state.get('resume_name', ''), "email": st.session_state.get('resume_email', ''),
                                       "phone": st.session_state.get('resume_phone', ''), "linkedin": st.session_state.get('resume_linkedin', ''),
                                       "github": st.session_state.get('resume_github', ''), "custom_links": st.session_state.custom_links,
                                       "location": st.session_state.get('resume_location', ''), "summary": st.session_state.get('resume_summary', ''),
                                       "skills": st.session_state.get('resume_skills', ''), "experience": st.session_state.experience,
                                       "education": st.session_state.education, "projects": st.session_state.projects,
                                       "achievements": st.session_state.achievements }
                    resume_text = format_resume_data_for_prompt(resume_data_qa)
                    if not resume_text: analysis_error = "❌ Error formatting data from Resume Maker."
                else: # Upload New Resume
                    resume_text = extract_text_from_file(uploaded_file_qa)
                    if not resume_text: analysis_error = "❌ Failed to extract text from the uploaded resume."
                if resume_text and not analysis_error:
                    formatted_prompt = predictive_qa_prompt_template.format(jd=jd_text_qa, resume_text=resume_text)
                    response_text = get_gemini_response(formatted_prompt, temperature=0.5)
                    if response_text.startswith("Error"): analysis_error = response_text
                    else: st.session_state.predictive_qa_output = response_text
            except Exception as e_process: analysis_error = f"An unexpected error occurred: {e_process}"
            finally: loading_placeholder_qa.empty()
            if analysis_error: st.session_state.predictive_qa_output = ""; st.error(analysis_error, icon="🚨")
            elif st.session_state.predictive_qa_output: st.success("Personalized questions generated!", icon="✅")
            
    if st.session_state.predictive_qa_output:
        st.divider(); st.subheader("💡 Your Personalized Interview Questions")
        st.info("Use these questions to prepare STAR method answers.")
        st.markdown(st.session_state.predictive_qa_output)
        
        # --- UPDATED: Predictive Q&A Download Buttons (DOCX + PDF - PARALLEL FIX) ---
        col1, col2 = st.columns(2)
        
        docx_buffer = create_markdown_docx(st.session_state.predictive_qa_output, title="Predictive Interview Q&A Report")
        with col1:
            if docx_buffer: st.download_button(label="⬇️ Download Questions (.docx)", data=docx_buffer, file_name="Personalized_Interview_Questions.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, key="download_qa_docx")
            else: st.button("Error DOCX", disabled=True, use_container_width=True, key="error_qa_docx")
        
        pdf_buffer = create_markdown_pdf(st.session_state.predictive_qa_output, title="Predictive Interview Q&A Report")
        with col2:
            if pdf_buffer: st.download_button(label="⬇️ Download Questions (.pdf)", data=pdf_buffer, file_name="Personalized_Interview_Questions.pdf", mime="application/pdf", use_container_width=True, key="download_qa_pdf")
            else: st.button("Error PDF", disabled=True, use_container_width=True, key="error_qa_pdf")


# --- AI Portfolio Generator View ---
elif active_view == "🌐 AI Portfolio Generator":
    st.header("🌐 AI Portfolio Website Generator")
    st.info("Generates a high-end portfolio website with animations.")
    st.subheader("1. Provide Your Information")
    resume_source_pf = st.radio("Data source:", ["Manual Entry", "Load from Resume Maker", "Upload & Parse Resume"],
                                 key="resume_source_pf", horizontal=True)
    uploaded_file_pf = None
    if resume_source_pf == "Upload & Parse Resume":
        uploaded_file_pf = st.file_uploader("Upload Resume for Parsing", type=["pdf", "docx"], key="uploader_pf_parse")
        st.info("ℹ️ AI parses this to fill fields. Separate from the PDF for the 'View Resume' button.")
        
    if resume_source_pf == "Load from Resume Maker":
        if st.button("Click to Load Data", key="load_pf_data"):
            st.session_state.pf_name = st.session_state.get('resume_name', '')
            st.session_state.pf_email = st.session_state.get('resume_email', '')
            st.session_state.pf_linkedin = st.session_state.get('resume_linkedin', '')
            st.session_state.pf_github = st.session_state.get('resume_github', '')
            st.session_state.pf_summary = st.session_state.get('resume_summary', '')
            st.session_state.pf_skills = st.session_state.get('resume_skills', '')
            st.session_state.pf_hero_text = ""
            st.session_state.portfolio_projects = []
            for proj in st.session_state.get('projects', []):
                link_found = ""; desc_lines = []
                for line in proj.get('description', '').split('\n'):
                    line_stripped = line.strip().lstrip('*- ')
                    if ("http://" in line_stripped or "https://" in line_stripped) and not link_found: link_found = line_stripped
                    elif line_stripped: desc_lines.append(line_stripped)
                st.session_state.portfolio_projects.append({"name": proj.get('name', ''), "description": " ".join(desc_lines), "link": link_found})
            st.success("Data loaded. Review fields below.", icon="✅"); st.rerun()
            
    if resume_source_pf == "Upload & Parse Resume" and uploaded_file_pf:
        if st.button("🤖 Parse Resume to Fill Fields", key="parse_resume_pf"):
            with st.spinner("🤖 AI is reading your resume..."):
                parse_error = None; text = extract_text_from_file(uploaded_file_pf)
                if text:
                    parse_prompt = resume_parser_for_portfolio_prompt.format(text=text)
                    response_text = get_gemini_response(parse_prompt, temperature=0.0)
                    try:
                        parsed_data = json.loads(clean_json_response(response_text))
                        st.session_state.pf_name = parsed_data.get('name', st.session_state.pf_name)
                        st.session_state.pf_email = parsed_data.get('email', st.session_state.pf_email)
                        st.session_state.pf_linkedin = parsed_data.get('linkedin', st.session_state.pf_linkedin)
                        st.session_state.pf_github = parsed_data.get('github', st.session_state.pf_github)
                        st.session_state.pf_summary = parsed_data.get('summary', st.session_state.pf_summary)
                        st.session_state.pf_skills = parsed_data.get('skills', st.session_state.pf_skills)
                        st.session_state.pf_hero_text = ""
                        st.session_state.portfolio_projects = parsed_data.get('projects', [])
                    except Exception as e_parse: parse_error = f"Error parsing AI response: {e_parse}\nRaw:{response_text}"
                else: parse_error = "❌ Failed to extract text from the resume."
            if parse_error: st.error(parse_error, icon="🚨")
            else: st.success("✅ Resume parsed! Review fields.", icon="✅"); st.rerun()
            
    st.divider()
    st.subheader("2. Review Your Portfolio Details")
    st.caption("Review/edit the fields below. This content will appear on your portfolio.")
    col1, col2 = st.columns(2)
    with col1:
        st.text_input("Full Name", key="pf_name"); st.text_input("Email", key="pf_email")
        st.text_area("About Me Summary", key="pf_summary", height=150)
    with col2:
        st.text_input("LinkedIn URL", key="pf_linkedin"); st.text_input("GitHub URL", key="pf_github")
        st.text_area("Skills (Comma-separated)", key="pf_skills", height=150)
    st.text_area("Typewriter Hero Text (Phrases separated by | )", placeholder="e.g., A Full-Stack Developer | A Creative Problem Solver", key="pf_hero_text")
    st.subheader("3. Portfolio Projects")
    st.info("Add/edit your projects. Include links!")
    with st.form("portfolio_project_form", clear_on_submit=True):
        proj_name = st.text_input("Project Name"); proj_desc = st.text_area("Description (1-2 lines)", height=100)
        proj_link = st.text_input("Project Link (URL)", placeholder="https://github.com/...")
        submitted_proj = st.form_submit_button("➕ Add Project")
        if submitted_proj:
            if proj_name and proj_desc and proj_link:
                st.session_state.portfolio_projects.append({"name": proj_name, "description": proj_desc, "link": proj_link})
                st.success("Project added!", icon="✅"); st.rerun()
            else: st.warning("Please fill all project fields, including the link.", icon="⚠️")
    if st.session_state.portfolio_projects:
        st.write("---"); st.write("**Current Projects:**")
        indices_to_remove = []
        for i, proj in enumerate(st.session_state.portfolio_projects):
            with st.container(border=True):
                st.markdown(f"**{i+1}. {proj.get('name', 'Unnamed Project')}**")
                st.caption(f"Desc: {proj.get('description', '')}"); st.caption(f"Link: {proj.get('link', '')}")
                if st.button("Remove 🗑️", key=f"rem_pf_proj_{i}", type="secondary"): indices_to_remove.append(i)
        if indices_to_remove:
             for index in sorted(indices_to_remove, reverse=True): del st.session_state.portfolio_projects[index]
             st.rerun()
             
    st.divider()
    st.subheader("4. Upload Your Resume File (Optional)")
    st.warning("Upload the final PDF here to enable the 'View Resume' button on your website.")
    uploaded_resume_file = st.file_uploader("Upload final Resume PDF", type="pdf", key="uploader_pf_resume_file")
    st.divider()
    
    st.subheader("5. Generate Your Website")
    if st.button("🚀 Generate AI Portfolio", type="primary", use_container_width=True, key="submit_portfolio"):
        st.session_state.portfolio_generated_code = ""; generation_error = None
        if not st.session_state.get("pf_name") or not st.session_state.get("pf_summary"): st.error("🚨 Please fill in at least Name and Summary.")
        elif not st.session_state.portfolio_projects: st.error("🚨 Please add at least one project.")
        else:
            loading_placeholder_pf = st.empty()
            lottie_data = None
            # Using a simpler Lottie filename for higher reliability
            try: lottie_data = load_lottiefile("STUDENT.lottie") 
            except Exception as e: st.warning(f"Could not load animation: {e}", icon="⚠️")
            if lottie_data:
                 with loading_placeholder_pf: st_lottie(lottie_data, height=300, key="lottie_pf_generate")
            else:
                 with loading_placeholder_pf: st.info("✨ Generating your portfolio website...")
            try:
                resume_data_uri = ""
                if uploaded_resume_file:
                    try:
                        resume_bytes = uploaded_resume_file.getvalue()
                        resume_base64 = base64.b64encode(resume_bytes).decode('utf-8')
                        resume_data_uri = f"data:application/pdf;base64,{resume_base64}"
                    except Exception as e_resume:
                        st.warning(f"⚠️ Resume file error: {e_resume}. Button omitted.", icon="⚠️")
                data_to_generate = {
                    "name": st.session_state.get("pf_name", ""), "email": st.session_state.get("pf_email", ""),
                    "linkedin": st.session_state.get("pf_linkedin", ""), "github": st.session_state.get("pf_github", ""),
                    "summary": st.session_state.get("pf_summary", ""),
                    "hero_text": st.session_state.get("pf_hero_text", "Developer | Problem Solver").split('|'),
                    "skills": [s.strip() for s in st.session_state.get("pf_skills", "").split(',') if s.strip()],
                    "projects": st.session_state.portfolio_projects }
                data_json = json.dumps(data_to_generate, indent=2)
                prompt = portfolio_generator_prompt.format(data_json=data_json, resume_data_uri=resume_data_uri)
                response = get_gemini_response(prompt, temperature=0.1)
                if response.startswith("Error"): generation_error = response
                else: st.session_state.portfolio_generated_code = response
            except Exception as e_process: generation_error = f"An error occurred: {e_process}"
            finally: loading_placeholder_pf.empty()
            if generation_error: st.session_state.portfolio_generated_code = ""; st.error(generation_error, icon="🚨")
            elif st.session_state.portfolio_generated_code: st.success("Website code generated!", icon="✅")
            
    if st.session_state.portfolio_generated_code:
        st.divider(); st.subheader("Your Generated Website Code")
        clean_code = clean_html_response(st.session_state.portfolio_generated_code)
        if clean_code:
            st.download_button(label="⬇️ Download index.html", data=clean_code, file_name="index.html", mime="text/html")
            with st.expander("View HTML code"): st.code(clean_code, language="html")
            st.subheader("Preview")
            st.info("Note: Previews may not be perfect. Download file for best result.")
            st.components.v1.html(clean_code, height=600, scrolling=True)
        else:
            st.error("Failed to extract clean HTML code from the AI response.", icon="🚨")
            st.text("Raw Response:"); st.code(st.session_state.portfolio_generated_code)

# ======================================================
# 🔗 LINKEDIN OPTIMIZER
# ======================================================
# --- LinkedIn Optimizer View ---
elif active_view == "🔗 LinkedIn Optimizer":
    st.header("🔗 LinkedIn Profile Optimizer")
    st.markdown("""
    **How to use:**
    1. Go to your LinkedIn Profile
    2. Click the **'More'** button in your intro section
    3. Select **'Save to PDF'**
    4. Upload that file here for analysis
    """)

    uploaded_ln = st.file_uploader("Upload LinkedIn PDF", type=["pdf"])

    if st.button("Analyze My Profile ✨", type="primary"):
        if not uploaded_ln:
            st.error("Please upload your LinkedIn PDF first.")
            st.stop()
            
        with st.spinner("🔍 Analyzing your LinkedIn profile..."):
            try:
                # Extract text from uploaded PDF
                text = extract_text_from_file(uploaded_ln)
                
                if not text or len(text.strip()) < 100:  # Basic validation
                    st.error("The uploaded PDF doesn't contain enough text. Please make sure it's a valid LinkedIn profile export.")
                    st.stop()
                
                # Prepare and send to Gemini API
                prompt = linkedin_optimization_prompt.format(resume_text=text[:15000])  # Limit text length
                
                # Get response from Gemini
                raw_response = get_gemini_response(prompt, temperature=0.3)
                
                # Clean and parse the response
                try:
                    # Clean up the response (handle markdown code blocks if present)
                    clean_response = clean_json_response(raw_response)
                    data = json.loads(clean_response)
                    
                    # Display results
                    st.success("✅ Analysis Complete!")
                    
                    # Create a two-column layout for scores
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.metric("Profile Completeness", data.get("Completeness_Score", "N/A"))
                    
                    with col2:
                        st.metric("Search Ranking", data.get("Search_Ranking_Score", "N/A"))
                    
                    # Display optimized headline
                    st.subheader("🚀 Optimized Headline")
                    st.info(data.get("Optimal_Headline", "No suggestion available"))
                    
                    # Display optimized summary
                    st.subheader("📝 Optimized 'About' Section")
                    st.write(data.get("Optimized_Summary", "No summary available"))
                    
                    # Display experience improvements
                    experience = data.get("Experience_Rewrite", [])
                    if experience:
                        with st.expander("💼 Experience Improvements (Click to Expand)", expanded=False):
                            for i, item in enumerate(experience, 1):
                                st.markdown(f"**Suggestion {i}**")
                                st.markdown(f"*Original:* {item.get('Original', 'N/A')}")
                                st.markdown(f"*Improved:* {item.get('Improved', 'N/A')}")
                                if i < len(experience):
                                    st.divider()
                    
                    # Display missing keywords
                    missing_keywords = data.get("Missing_Profile_Keywords", [])
                    if missing_keywords:
                        with st.expander("🔍 Missing Keywords (Click to Expand)", expanded=False):
                            st.markdown("These keywords could improve your profile's search visibility:")
                            st.write(", ".join([f"`{kw}`" for kw in missing_keywords]))
                    
                    # Display keyword density analysis
                    keyword_density = data.get("Keyword_Density", {})
                    if keyword_density:
                        with st.expander("📊 Keyword Analysis (Click to Expand)", expanded=False):
                            st.markdown("**Strong Keywords:** " + ", ".join([f"`{kw}`" for kw in keyword_density.get("Strong", [])]))
                            st.markdown("**Could Be Stronger:** " + ", ".join([f"`{kw}`" for kw in keyword_density.get("Weak", [])]))
                            st.markdown("**Missing Keywords:** " + ", ".join([f"`{kw}`" for kw in keyword_density.get("Missing", [])]))
                    
                    # Display behavioral strategy
                    strategy = data.get("Behavioral_Strategy", {})
                    if strategy:
                        with st.expander("🎯 Engagement Strategy (Click to Expand)", expanded=False):
                            st.markdown("**Visual Showcase Suggestions:**")
                            for tip in strategy.get("Visual_Showcase_Suggestions", []):
                                st.markdown(f"- {tip}")
                            
                            st.markdown("\n**Call to Action Strategy:**")
                            st.write(strategy.get("Call_to_Action_Strategy", "Not provided"))
                    
                    # Create a download button for the full report
                    report = {
                        "LinkedIn Optimization Report": {
                            "date": datetime.datetime.now().strftime("%Y-%m-%d"),
                            "completeness_score": data.get("Completeness_Score"),
                            "search_ranking_score": data.get("Search_Ranking_Score"),
                            "optimal_headline": data.get("Optimal_Headline"),
                            "optimized_summary": data.get("Optimized_Summary"),
                            "experience_improvements": data.get("Experience_Rewrite", []),
                            "missing_keywords": data.get("Missing_Profile_Keywords", []),
                            "keyword_density": data.get("Keyword_Density", {}),
                            "behavioral_strategy": data.get("Behavioral_Strategy", {})
                        }
                    }
                    
                    st.download_button(
                        label="📥 Download Full Report (JSON)",
                        data=json.dumps(report, indent=2, ensure_ascii=False),
                        file_name=f"linkedin_optimization_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.json",
                        mime="application/json"
                    )
                    
                except json.JSONDecodeError:
                    st.error("Failed to parse the optimization response. The AI might have returned an invalid format.")
                    st.text_area("Raw Response", raw_response, height=300)
                
            except Exception as e:
                st.error(f"An error occurred while processing your request: {str(e)}")
                st.exception(e)

# ======================================================
# 🔍 JOB SEARCH — ADVANCED (RapidAPI JSearch)
# ======================================================
elif active_view == "🔍 Job Search":

    # Initialize session storage
    if "job_results" not in st.session_state:
        st.session_state.job_results = []
    if "job_page" not in st.session_state:
        st.session_state.job_page = 0

    st.title("🔍 AI Job Search")

    # Inputs
    query = st.text_input("Job Title or Keywords", value="")
    location = st.text_input("Location (optional)", value="")

    experience_range = st.selectbox(
        "Experience Level",
        ["Any", "Fresher (0 years)", "0–1 years", "1–3 years", "3–5 years", "5–8 years", "8+ years"]
    )
    mapped_experience = map_experience_to_api(experience_range)

    job_type = st.selectbox("Job Type", ["Any", "Full-time", "Part-time", "Contract", "Internship"])
    posted = st.selectbox("Posted Within", ["all", "today", "3days", "week", "month"])

    # How many pages to fetch (max_pages). Keep reasonable default (10)
    max_pages = 10

    # Search action
    if st.button("Search Jobs 🔎", type="primary"):
        if not query.strip():
            st.error("Enter a job title first.")
            st.stop()

        normalized_location = normalize_location(location)
        final_query = build_job_query(query, normalized_location, experience_range, job_type)

        with st.spinner("Fetching jobs…"):
            api_results = search_jobs_rapidapi(
                query=final_query,
                location=normalized_location,
                experience=mapped_experience,
                job_type=job_type,
                posted=posted,
                max_pages=max_pages
            )

        # store results
        st.session_state.job_results = api_results
        st.session_state.job_page = 0

    # Initialize job_results in session state if it doesn't exist
    if 'job_results' not in st.session_state:
        st.session_state.job_results = []
    if 'job_page' not in st.session_state:
        st.session_state.job_page = 0
    
    # Get results from session
    results = st.session_state.job_results
    if not results:
        st.info("No results yet — run a search to fetch jobs.")
        st.stop()

    # ============================================
    # PAGINATION CONFIG
    page_size = 5
    total = len(results)
    total_pages = max(1, (total + page_size - 1) // page_size)
    
    # Ensure page is within bounds
    page = st.session_state.job_page
    if page >= total_pages:
        page = max(0, total_pages - 1)
        st.session_state.job_page = page
        st.rerun()

    # Get current page of results
    start = page * page_size
    end = min((page + 1) * page_size, total)
    current_jobs = results[start:end]

    st.success(f"Showing jobs {start+1} to {end} of {total} — Page {page+1}/{total_pages}")
    
    # Jump to page control
    sel_page = st.selectbox(
        "Jump to Page",
        list(range(1, total_pages + 1)),
        index=page,
        key="jump_page_selector"
    )

    if int(sel_page) - 1 != st.session_state.job_page:
        st.session_state.job_page = int(sel_page) - 1
        st.rerun()

    # Previous/Next buttons
    prev_col, _, next_col = st.columns([1, 5, 1])
    with prev_col:
        if st.button("⬅ Previous", key="prev_btn"):
            if st.session_state.job_page > 0:
                st.session_state.job_page -= 1
                st.rerun()
    with next_col:
        if st.button("Next ➡", key="next_btn"):
            if st.session_state.job_page < total_pages - 1:
                st.session_state.job_page += 1
                st.rerun()

    # Display current page of jobs
    for job in current_jobs:
        st.divider()
        st.subheader(job.get("job_title", "No Title"))
        st.write("**Company:**", job.get("employer_name", "N/A"))
        st.write("**Location:**", job.get("job_city", ""), job.get("job_country", ""))
   

    # Salary + INR conversion
    salary_display = extract_salary(job)

    try:
        if isinstance(salary_display, str) and "–" in salary_display:
            parts = salary_display.split("–")
            lo = int(parts[0].strip().replace(",", ""))
            hi = int(parts[1].split("/")[0].strip().replace(",", ""))

            lo_inr = convert_to_inr(lo)
            hi_inr = convert_to_inr(hi)

            salary_display += f"  |  🇮🇳 ₹{lo_inr:,} – ₹{hi_inr:,}"
    except:
        pass

    st.write("**Salary / Package:**", salary_display)

    # Description
    desc = job.get("job_description", "No description available")
    st.write(desc[:600] + "...")

    # Apply link
    apply_link = job.get("job_apply_link") or job.get("job_apply_url") or job.get("job_apply")
    if apply_link:
        st.markdown(f"[Apply Now 🚀]({apply_link})", unsafe_allow_html=True)

    # ============================================
    # SALARY BENCHMARK BUTTON (Stable Unique Key)
    # ============================================
    job_title = job.get("job_title", "")
    city = job.get("job_city", "") or normalized_location or "India"

    uid = f"{page}_{global_idx}"
    sb_key = f"sb_btn_{uid}"

    if st.button(f"💰 Salary Benchmark for {job_title}", key=sb_key):

        st.subheader("💰 Salary Benchmark Report")

        # Fetch salary API
        with st.spinner("Fetching salary insights…"):
            sb = get_salary_benchmark(job_title, city)

        if sb.get("fallback"):
            st.warning("Salary API unavailable — using AI estimate only.")
            salary_min = salary_max = None
        else:
            salary_min = sb.get("salary_min", 0)
            salary_max = sb.get("salary_max", 0)

            if salary_min and salary_max:
                st.write(f"### 📌 Market Range in {city}")
                st.success(f"₹{salary_min:,} – ₹{salary_max:,}")

        # AI resume evaluation
        resume_text = st.session_state.get("global_resume_text", "")
        if not resume_text:
            st.info("Upload a resume in Resume Analyzer to enable AI-based salary assessment.")
        else:
            with st.spinner("Analyzing your resume…"):
                ai_val = ai_market_value(job_title, city, resume_text, salary_min, salary_max)

            st.write("### 📌 Your Resume Market Value")
            st.success(ai_val.get("resume_market_value", "N/A"))

            st.write("### 🎯 Expected Range")
            st.info(ai_val.get("expected_range", "N/A"))

            st.write("### 💬 Negotiation Tip")
            st.write(ai_val.get("negotiation_tip", "N/A"))
